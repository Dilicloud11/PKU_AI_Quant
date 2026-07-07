# -*- coding: utf-8 -*-
"""
生成 TASK3 作业文档：张哲铭TASK3.docx / .pdf
格式：宋体、五号(10.5pt)、1.5倍行距、0段间距、正文两端对齐。
内容：双均线策略原理、量化评估指标、编程实现、文献综述、
      多标的×多周期回测、超额收益与风险分析、应用心得。
作者：张哲铭
"""
import os
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

BASE = os.path.dirname(os.path.abspath(__file__))
FIG = os.path.join(BASE, "figures")
DOCX = os.path.join(BASE, "张哲铭TASK3.docx")
FONT = "宋体"
SIZE = Pt(10.5)  # 五号

# 标的名映射
NAME = {
    "sh600900": "长江电力", "hk00700": "腾讯控股", "sh518880": "黄金ETF",
    "sh515450": "红利低波50ETF", "sz159941": "纳指ETF", "sh588000": "科创50ETF",
    "sh510300": "沪深300ETF", "sh510500": "中证500ETF",
}


# ============ 样式辅助 ============
def rfont(run, size=SIZE, bold=False, color=None, name=FONT):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = size
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def pfmt(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.alignment = align


def body(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False, indent=True, size=SIZE):
    p = doc.add_paragraph(); pfmt(p, align)
    if indent:
        p.paragraph_format.first_line_indent = Pt(21)
    rfont(p.add_run(text), size=size, bold=bold)
    return p


def title(doc, text, size=Pt(16)):
    p = doc.add_paragraph(); pfmt(p, WD_ALIGN_PARAGRAPH.CENTER)
    rfont(p.add_run(text), size=size, bold=True); return p


def heading(doc, text, size=Pt(13)):
    p = doc.add_paragraph(); pfmt(p, WD_ALIGN_PARAGRAPH.LEFT)
    rfont(p.add_run(text), size=size, bold=True); return p


def subheading(doc, text, size=Pt(11)):
    p = doc.add_paragraph(); pfmt(p, WD_ALIGN_PARAGRAPH.LEFT)
    p.paragraph_format.first_line_indent = Pt(21)
    rfont(p.add_run(text), size=size, bold=True); return p


def caption(doc, text):
    p = doc.add_paragraph(); pfmt(p, WD_ALIGN_PARAGRAPH.CENTER)
    rfont(p.add_run(text), size=Pt(9), bold=True); return p


def pic(doc, name, width=6.1):
    p = doc.add_paragraph(); pfmt(p, WD_ALIGN_PARAGRAPH.CENTER)
    p.add_run().add_picture(os.path.join(FIG, name), width=Inches(width))


def formula(doc, text):
    p = doc.add_paragraph(); pfmt(p, WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run(text)
    r.font.name = "Cambria Math"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(10.5); r.font.italic = True


def code_block(doc, code):
    for line in code.split("\n"):
        p = doc.add_paragraph(); pfmt(p, WD_ALIGN_PARAGRAPH.LEFT)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        r = p.add_run(line if line else " ")
        r.font.name = "Consolas"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def make_table(doc, headers, rows, widths=None, size=Pt(8.5)):
    table = doc.add_table(rows=1, cols=len(headers)); table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        rfont(c.paragraphs[0].add_run(h), size=Pt(8.5), bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = cells[i].paragraphs[0].paragraph_format
            pf.space_before = Pt(0); pf.space_after = Pt(0)
            rfont(cells[i].paragraphs[0].add_run(str(v)), size=size)
    return table


def pct(x, d=1):
    return f"{x*100:.{d}f}%"


def build():
    df = pd.read_csv(os.path.join(BASE, "backtest_results.csv"))

    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    st.font.size = SIZE

    # ===== 封面标题 =====
    title(doc, "策略首秀：用均线交叉反应市场趋势变化")
    info = doc.add_paragraph(); pfmt(info, WD_ALIGN_PARAGRAPH.CENTER)
    rfont(info.add_run("北京大学 AI 量化工作坊 · TASK3　　姓名：张哲铭"), size=Pt(10.5))
    doc.add_paragraph()

    body(doc,
        "本任务在前两次工作坊（搭建数据引擎、数据诊断与指标构造）的基础上，迈出策略开发的第一步："
        "系统学习并实现经典的双均线（Dual Moving Average）交叉策略，用金叉、死叉信号捕捉市场趋势变化，"
        "并以最大回撤、夏普比率、累计回报、超额收益、胜率、盈亏比等指标全面评估策略表现。为使结论更可靠，"
        "本文先检索了均线策略的重要学术文献，明确其在何种标的上更有效，随后对黄金 ETF、腾讯控股、纳指 ETF、"
        "科创 50、沪深 300、中证 500 等 8 个标的、短/中/长三组均线周期共 24 组参数进行了实证回测。")

    # ===== 一、双均线策略原理 =====
    heading(doc, "一、双均线策略：金叉与死叉")
    body(doc,
        "移动平均线（Moving Average, MA）是把最近 N 个交易日的收盘价取算术平均得到的一条平滑曲线，"
        "它过滤了单日价格的随机噪音，反映一段时间内价格的“平均成本”与趋势方向。均线周期越短，对价格变化越"
        "敏感、越贴近行情；周期越长，越平滑、越能代表中长期趋势。")
    body(doc,
        "双均线策略同时使用两条不同周期的均线：一条较短（如 MA5、MA20，称“快线”），一条较长"
        "（如 MA20、MA60，称“慢线”）。二者的相对位置刻画了短期动能与长期趋势的关系，其交叉点即为交易信号：")
    subheading(doc, "· 金叉（Golden Cross）——买入信号")
    body(doc,
        "当短期均线由下向上穿越长期均线时，形成“金叉”。它意味着近期价格上涨的动能已强于中长期均值，"
        "市场趋势可能由弱转强、由跌转涨，是经典的做多（买入）信号。")
    subheading(doc, "· 死叉（Death Cross）——卖出信号")
    body(doc,
        "当短期均线由上向下穿越长期均线时，形成“死叉”。它意味着近期动能已弱于中长期均值，趋势可能"
        "由强转弱、由涨转跌，是经典的做空或平仓（卖出）信号。")
    body(doc,
        "双均线策略的本质是一种“趋势跟踪”：金叉后满仓持有、死叉后清仓空仓，力求“截断亏损、让利润奔跑”，"
        "在单边上涨行情中吃到主升浪、在下跌行情中及时离场规避风险。它不预测顶底，只对已经发生的趋势变化做出反应，"
        "因而天然存在“滞后性”——这也是后文实证中需要重点评估的代价。")

    # ===== 二、量化评估指标 =====
    heading(doc, "二、量化策略评估的核心指标")
    body(doc,
        "评价一个策略不能只看“赚了多少”，还要看“冒了多大风险、是否稳定、是否真的比躺着不动更强”。"
        "本文采用以下一整套指标进行评估：")

    subheading(doc, "1. 累计回报（Cumulative Return）与总收益率")
    body(doc,
        "指从期初到期末，策略净值累计增长的比例。若初始净值为 1、期末净值为 V，则累计回报 = V − 1。"
        "它是最直观的盈利结果，但没有考虑时间长度与风险。")
    formula(doc, "净值曲线：Equity_t = Π(1 + r_i)，  累计回报 = Equity_末 − 1")

    subheading(doc, "2. 年化收益率（Annualized Return）")
    body(doc,
        "把不同时间长度的收益折算到“每年”的标准口径，便于横向比较。以 252 个交易日为一年：")
    formula(doc, "年化收益率 = (1 + 累计回报) ^ (252 / 交易天数) − 1")

    subheading(doc, "3. 最大回撤（Maximum Drawdown, MDD）")
    body(doc,
        "指净值曲线从历史最高点回落到之后最低点的最大跌幅，衡量策略“最坏情况下会亏多少”，是刻画"
        "下行风险与持有煎熬程度的关键指标。MDD 越小（越接近 0），资金曲线越稳、回撤控制越好。")
    formula(doc, "MDD = min_t ( Equity_t / max_{s≤t} Equity_s − 1 )")

    subheading(doc, "4. 夏普比率（Sharpe Ratio）")
    body(doc,
        "由诺贝尔经济学奖得主 William Sharpe 提出，衡量“每承担一单位波动风险，能换来多少超额收益”，"
        "是风险调整后收益的核心标尺。数值越高越好，一般认为大于 1 即较为优秀。计算时以无风险利率为基准"
        "（教学场景取 0），用日收益的均值除以标准差再年化：")
    formula(doc, "夏普比率 = (策略日收益均值 − 无风险日利率) / 策略日收益标准差 × √252")

    subheading(doc, "5. 超额收益（Excess Return）——策略是否真的“有本事”")
    body(doc,
        "这是本文特别强调的指标。一只标的本身可能就是大牛股，即使什么都不做、一直满仓持有也能赚很多。"
        "为了区分“收益到底是策略择时带来的，还是标的本身优秀带来的”，本文定义：")
    formula(doc, "超额收益 = 策略年化收益 − 买入持有（Buy&Hold）年化收益")
    body(doc,
        "超额收益为正，说明策略的择时确实创造了价值、跑赢了“躺平”；为负，则说明频繁进出反而不如一直持有。"
        "它把“标的 β”从“策略 α”中剥离出来，是判断策略真实能力最诚实的一把尺子。")

    subheading(doc, "6. 胜率（Win Rate）与盈亏比（Profit/Loss Ratio）")
    body(doc,
        "胜率 = 盈利交易笔数 / 总交易笔数，反映“做对的概率”；盈亏比 = 平均每笔盈利 / 平均每笔亏损，"
        "反映“做对时赚得多、做错时亏得少”的程度。趋势跟踪策略的典型特征是“胜率不高但盈亏比很高”——"
        "靠少数几次抓住大趋势的盈利，覆盖多次小幅止损，这一点在后文黄金 ETF 的结果中体现得非常明显。")

    # ===== 三、Python 编程实现 =====
    heading(doc, "三、Python 编程实现")
    body(doc,
        "本文的回测引擎完全手工实现（strategy.py），核心分为五步：加载数据、计算双均线、生成金叉死叉信号、"
        "按“信号次日生效”建立仓位、计入交易成本后累乘得到净值曲线。为杜绝“未来函数”（用当天才知道的信息去"
        "交易当天），第 t 日收盘后产生的信号在第 t+1 日才实际建仓或平仓，通过 position = signal.shift(1) 实现；"
        "同时在仓位发生变化的当日扣除单边万分之五的交易成本。核心代码如下：")
    code_block(doc,
        '# 1) 计算短、长均线\n'
        'ma_short = close.rolling(short).mean()\n'
        'ma_long  = close.rolling(long).mean()\n\n'
        '# 2) 金叉/死叉信号：短均线在长均线之上则持仓状态=1\n'
        'raw = (ma_short > ma_long).astype(int)\n'
        'cross = raw.diff()          # +1 金叉当日, -1 死叉当日\n\n'
        '# 3) 信号次日生效，避免未来函数\n'
        'position = raw.shift(1).fillna(0)\n\n'
        '# 4) 计入交易成本后的策略日收益\n'
        'ret = close.pct_change()\n'
        'trade = position.diff().abs()          # 仓位变化 -> 发生交易\n'
        'strat_ret = position * ret - trade * 0.0005\n\n'
        '# 5) 净值曲线（策略 vs 买入持有基准）\n'
        'equity       = (1 + strat_ret).cumprod()\n'
        'bench_equity = (1 + ret).cumprod()')
    body(doc,
        "数据方面，本文通过腾讯自选股行情接口获取 8 个标的近 8 年（多数为 2018-04 至 2026-07，共约 2000 个"
        "交易日）的前复权日线数据。前复权可消除分红、拆股对价格的跳变影响，使均线计算与收益回测更准确。"
        "下面以黄金 ETF（518880）的中线参数（MA20×MA60）为例，展示价格、双均线与买卖信号的可视化。")
    pic(doc, "signal_sh518880.png")
    caption(doc, "图 1　黄金 ETF（518880）双均线交易信号（MA20×MA60）")
    body(doc,
        "如图 1，红色上三角为金叉买入点、绿色下三角为死叉卖出点。可以清晰看到：在 2018—2019 与 2021—2023 的"
        "横盘震荡期，均线频繁交叉、产生了较多“来回打脸”的假信号；而在 2019—2020 与 2024—2026 的两轮单边上涨中，"
        "金叉后价格持续走高、策略得以长期持有并吃到主升浪，直到 2026 年中价格见顶回落时死叉离场。这直观展示了"
        "双均线策略“怕震荡、爱趋势”的本质特征。")

    # ===== 四、文献综述：均线策略在什么标的上更有效 =====
    heading(doc, "四、文献综述：双均线策略在什么标的上验证更有效")
    body(doc,
        "在动手回测前，先回顾学术界对均线/趋势跟踪策略的重要研究，以明确“该策略在什么标的上更靠谱”。")
    subheading(doc, "1. 奠基之作：Brock, Lakonishok & LeBaron (1992)")
    body(doc,
        "发表于顶级期刊《Journal of Finance》的 Brock、Lakonishok 与 LeBaron（1992，简称 BLL）是检验技术交易"
        "规则的开山之作。他们用道琼斯工业指数 1897—1986 长达 90 年的数据，对移动平均规则和区间突破规则进行"
        "了严格的统计（自助法 Bootstrap）检验，发现：买入信号（短均线上穿长均线）之后的平均收益显著高于卖出"
        "信号之后的收益，且买入信号后的收益波动更小。这一结果无法被随机游走、AR(1)、GARCH-M 等主流模型解释，"
        "为“简单均线规则确实包含预测力”提供了首个量化证据，动摇了有效市场假说。")
    subheading(doc, "2. 有效性会衰减：自适应市场假说（Adaptive Market Hypothesis）")
    body(doc,
        "然而后续研究（如 Bessembinder & Chan 1998；以及对 1987—2013 年美、英、日三大成熟市场的再检验）发现，"
        "BLL 记录的均线规则预测力在 1987 年之后显著减弱。原因在于：随着市场参与者学习、套利资金涌入、交易成本"
        "下降，一旦某个规则被广泛知晓，其超额收益就会被“抢跑”而逐渐消失。这正是 Lo（2004）提出的“自适应市场"
        "假说”——技术规则只在一段时间内有效，会随市场进化而失灵。这提醒我们：在流动性极好、参与者极成熟的"
        "大盘蓝筹/大盘指数上，均线择时越来越难跑赢买入持有。")
    subheading(doc, "3. 大宗商品与趋势性资产更适合：Miffre & Rallis (2007) 等")
    body(doc,
        "与成熟股市形成对比，趋势跟踪类策略在大宗商品期货上表现突出。Miffre 与 Rallis（2007，发表于《Journal of "
        "Banking & Finance》）对 31 种商品期货 1979—2004 的动量策略研究发现，13 个动量策略年均收益达 9.38%，"
        "而同期等权买入持有商品组合反而亏损 2.64%；且这一收益不因样本期而衰减、与股债相关性低。Moskowitz、Ooi "
        "与 Pedersen（2012）在 58 个跨资产品种上进一步证实了“时间序列动量”的普遍存在。其共同结论是：趋势跟踪"
        "策略在趋势性强、波动大、存在持续供求/宏观驱动的标的（如黄金等贵金属、原油等商品，以及具备长期上行"
        "beta 的成长型指数）上更为有效。")
    subheading(doc, "4. 对本文标的选择的启示")
    body(doc,
        "综合文献，双均线策略更可能在“黄金等大宗商品、趋势鲜明的成长型指数”上跑出价值，而在“高效率、宽幅"
        "震荡的大盘蓝筹与宽基指数”上容易因滞后与假信号而跑输买入持有。恰好，本文用户关注的标的中即包含"
        "黄金 ETF（518880，商品属性）与纳指 ETF（159941，成长指数），可直接作为文献结论的实证检验对象；"
        "同时纳入沪深 300、中证 500 等宽基指数作为对照。下一节即用真实回测数据加以验证。")

    # ===== 五、多标的 × 多周期回测结果 =====
    heading(doc, "五、多标的、多周期实证回测")
    body(doc,
        "本文对 8 个标的分别采用短线（MA5×MA20）、中线（MA20×MA60）、长线（MA60×MA120）三组均线参数进行回测，"
        "共 24 组结果。为回答“收益是策略带来的还是标的带来的”，每组均与“买入持有”基准对比，重点看超额收益。")

    # ---- 表1：论文验证标的（黄金ETF + 纳指ETF）三周期结果 ----
    subheading(doc, "1. 文献验证：黄金 ETF 与纳指 ETF")
    body(doc,
        "先看文献指向的两个标的。表 1 给出黄金 ETF 与纳指 ETF 在三组周期下的完整指标。")
    rows1 = []
    for code in ["sh518880", "sz159941"]:
        for lb in ["短线", "中线", "长线"]:
            r = df[(df["code"] == code) & (df["period_label"] == lb)].iloc[0]
            rows1.append([
                NAME[code], lb, f"MA{r['short']}×MA{r['long']}",
                pct(r["strat_total"]), pct(r["strat_annual"]),
                pct(r["excess_annual"]), f"{r['sharpe']:.2f}",
                pct(r["mdd"]), pct(r["win_rate"], 0), f"{r['pl_ratio']:.2f}",
            ])
    caption(doc, "表 1　黄金 ETF 与纳指 ETF 三组周期回测结果")
    make_table(doc,
        ["标的", "周期", "参数", "总收益", "年化", "超额年化", "夏普", "最大回撤", "胜率", "盈亏比"],
        rows1)
    body(doc,
        "结果与文献高度吻合：黄金 ETF 是全样本中双均线策略表现最好的标的——中线参数下年化 15.0%、夏普 1.02、"
        "最大回撤仅 24.9%，超额年化仅 -0.7%（几乎追平买入持有），盈亏比高达 10.7（胜率 64%）。这正是趋势跟踪"
        "的理想画像：黄金 2024 年以来的强单边上涨让策略长期满仓吃到主升浪，而 2026 年中的高位死叉又帮它规避了"
        "后续回撤。纳指 ETF 作为成长型指数同样表现不俗（短线夏普 0.78）。二者共同印证：趋势性强的商品与成长"
        "资产确实是双均线策略的“主场”。")

    # ---- 图2：黄金ETF净值曲线 ----
    pic(doc, "equity_sh518880.png")
    caption(doc, "图 2　黄金 ETF 三组周期策略净值 vs 买入持有")
    body(doc,
        "如图 2，黄金 ETF 的中长线策略净值在 2026 年高点后走平（死叉离场），成功躲过了买入持有（灰色虚线）从"
        "4.4 跌回 3.0 的剧烈回撤——这正是趋势跟踪“放弃部分顶部收益、换取回撤保护”价值的最佳写照。")

    # ---- 表2：用户关注标的中线结果全览 ----
    subheading(doc, "2. 用户关注标的：中线参数横向对比")
    body(doc,
        "表 2 汇总全部 8 个标的在中线参数（MA20×MA60）下的表现，并按超额年化收益从高到低排序。")
    mid = df[df["period_label"] == "中线"].copy().sort_values("excess_annual", ascending=False)
    rows2 = []
    for _, r in mid.iterrows():
        rows2.append([
            r["name"], pct(r["strat_annual"]), pct(r["bench_annual"]),
            pct(r["excess_annual"]), f"{r['sharpe']:.2f}",
            pct(r["mdd"]), pct(r["bench_mdd"]), f"{int(r['n_trades'])}",
            pct(r["win_rate"], 0),
        ])
    caption(doc, "表 2　8 标的中线策略（MA20×MA60）指标全览（按超额年化排序）")
    make_table(doc,
        ["标的", "策略年化", "基准年化", "超额年化", "夏普", "策略回撤", "基准回撤", "交易次数", "胜率"],
        rows2)
    pic(doc, "summary_excess.png")
    caption(doc, "图 3　各标的中线策略超额年化收益对比")
    body(doc,
        "如表 2 与图 3，一个诚实但重要的结论浮现：在 2018—2026 这段整体向上的样本里，几乎所有标的的双均线择时都"
        "跑输了买入持有（超额年化为负）。这与自适应市场假说一致——在流动性好的宽基指数（沪深 300、中证 500）与"
        "低波动标的（红利低波、长江电力）上，均线的滞后性让它频繁高买低卖，超额收益垫底（纳指 -12.9%、红利低波 "
        "-11.7%）。而黄金 ETF（-0.7%）与腾讯控股、中证 500 短线等趋势/波动更强的组合，超额损失最小甚至转正，"
        "再次验证了“标的选择比参数更重要”。")

    # ---- 图4：回撤对比 ----
    subheading(doc, "3. 别只看收益：双均线真正的价值在“控回撤”")
    body(doc,
        "若仅以超额收益论英雄，容易误判策略毫无价值。但把视角切换到风险维度，结论就大不相同。图 4 对比了中线"
        "策略与买入持有的最大回撤。")
    pic(doc, "risk_compare.png")
    caption(doc, "图 4　双均线策略 vs 买入持有：最大回撤对比（中线）")
    body(doc,
        "如图 4，在高波动标的上，双均线策略显著削减了回撤：腾讯控股从买入持有的 76.7% 大幅压降到 53.3%，"
        "科创 50 从 59.9% 降到 45.6%。也就是说，策略用“少赚一点”换来了“回撤浅很多、持有体验平稳很多”。"
        "反之，在本就低波动的红利低波、长江电力上，择时失误反而略微加大了回撤——这说明低波动标的没有明显趋势"
        "可供跟踪，不适合用双均线。")

    # ---- 图5：夏普热力图 ----
    pic(doc, "sharpe_compare.png")
    caption(doc, "图 5　8 标的 × 3 周期 夏普比率热力图")
    body(doc,
        "图 5 的夏普比率（风险调整后收益）进一步佐证：黄金 ETF 在三组周期上全部“翻绿”（0.89—1.02，全样本最高），"
        "纳指 ETF 次之，是最适合双均线策略的两类标的；而沪深 300、长江电力、红利低波等则大面积偏橙偏红，"
        "风险调整后性价比不佳。")

    # ---- 图6：收益-回撤散点，讨论周期效应 ----
    subheading(doc, "4. 短线、中线、长线：周期怎么选")
    body(doc,
        "把 24 组结果画到“收益—回撤”平面上（图 6），可以观察均线周期的整体规律。")
    pic(doc, "period_effect.png")
    caption(doc, "图 6　收益—回撤分布（8 标的 × 3 周期，共 24 组）")
    body(doc,
        "综合表 1 与图 6 可总结出周期效应：（1）短线（MA5×MA20）信号最灵敏、交易最频繁（单标的 38—70 次），"
        "在强趋势标的上能更早上车（如科创 50 短线年化 13.0%、超额转正 +5.8%），但在震荡标的上假信号最多、"
        "交易成本侵蚀严重（沪深 300 短线年化 -1.4%）。（2）长线（MA60×MA120）最稳健、交易最少（5—11 次）、"
        "盈亏比最高（黄金长线 10.5、腾讯长线胜率 80%），滞后性也最强、容易错过拐点。（3）中线（MA20×MA60）"
        "在灵敏与稳健之间较为均衡，是多数标的的折中之选。总体上，没有“万能周期”，需与标的波动特性匹配："
        "波动大、趋势急的标的适合偏短周期抢反应，波动温和、趋势慢的标的适合偏长周期滤噪音。")

    # ===== 六、应用心得 =====
    heading(doc, "六、双均线策略适用场景与应用心得")
    body(doc,
        "综合文献与本文 24 组实证，对双均线策略的适用场景与使用要点总结如下：")
    body(doc,
        "1. 适合“有趋势”的标的，最怕“宽幅震荡”。策略的全部收益来自抓住单边趋势，因此在黄金等大宗商品、"
        "纳指等趋势鲜明的成长指数上最有效（黄金 ETF 夏普高达 1.02）；而在长期横盘或高效率的宽基指数上，"
        "均线来回穿插产生大量假信号，反而不断“高买低卖”。标的选择的重要性远高于参数调优。", indent=True)
    body(doc,
        "2. 它的核心价值往往不是“多赚”，而是“少亏、控回撤”。在长牛样本里策略超额收益多为负，但它把腾讯的"
        "最大回撤从 77% 压到 53%、把黄金的回撤从 30% 压到 25%。对追求平稳、厌恶大幅回撤的资金而言，"
        "用一定的收益换取显著更浅的回撤与更高的夏普，本身就是有意义的风险管理。", indent=True)
    body(doc,
        "3. 天生“低胜率、高盈亏比”，要接受频繁小亏。本文趋势型标的胜率多在 40%—65%，但盈亏比常达 2—10 倍，"
        "盈利高度依赖少数几笔大趋势。使用者必须有纪律地执行每一次止损信号，不能因连续小亏而放弃，"
        "否则会错过真正贡献收益的那几笔大行情。", indent=True)
    body(doc,
        "4. 周期需与标的波动匹配，且必须计入交易成本、规避未来函数。波动大选偏短周期、波动小选偏长周期；"
        "短周期交易频繁，务必把佣金与冲击成本纳入回测（本文用万分之五），并让信号“次日生效”，"
        "否则回测会系统性高估收益。", indent=True)
    body(doc,
        "5. 单一均线信号并非万能，宜作为“趋势过滤器”与其他工具配合。文献显示其超额收益会随市场成熟而衰减，"
        "实务中常与 RSI/MACD（判断超买超卖与背离）、成交量、止损止盈规则、以及仓位管理结合，"
        "用震荡指标过滤震荡市假信号，用趋势指标确认方向，才能扬长避短。", indent=True)
    body(doc,
        "总之，双均线策略是理解“趋势跟踪”思想最好的入门范式：逻辑简单、可解释性强、能清晰暴露择时策略"
        "“怕震荡、爱趋势、重风控”的一般规律。它教会我们的最重要一课或许是——在评价任何策略时，都要用超额收益"
        "把“标的的 β”与“策略的 α”分开，用夏普和最大回撤把“收益”与“风险”一起看，才能做出诚实而全面的判断。")

    # 结尾：数据与代码说明
    doc.add_paragraph()
    note = doc.add_paragraph(); pfmt(note, WD_ALIGN_PARAGRAPH.LEFT)
    rfont(note.add_run("数据来源：腾讯自选股行情接口（前复权日线，2018-2026）；"
                       "回测与绘图代码：strategy.py / run_backtest.py / extra_analysis.py。"
                       "本文仅为量化学习实践，不构成任何投资建议，市场有风险，决策需谨慎。"),
          size=Pt(9), color=RGBColor(0x66, 0x66, 0x66))

    doc.save(DOCX)
    print("全部章节已写入：", DOCX)
    return doc, df


if __name__ == "__main__":
    build()
