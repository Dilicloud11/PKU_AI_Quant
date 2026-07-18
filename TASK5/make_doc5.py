# -*- coding: utf-8 -*-
"""
生成 TASK5 作业文档：张哲铭TASK5.docx / .pdf
格式：宋体、五号(10.5pt)、1.5 倍行距、0 段间距、正文两端对齐。
内容：分类型机器学习算法解释、模型评价指标（混淆矩阵/AUC/ROC）解释、
      重要文献综述、数据与特征工程、时间序列划分、多算法建模与网格调优、
      分类与回归评估指标对比、乳腺癌教学示例、全套 EDA 图表解读、结论。
作者：张哲铭
"""
import os
import json
import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

BASE = os.path.dirname(os.path.abspath(__file__))
FIG = os.path.join(BASE, "figures")
DATA = os.path.join(BASE, "data")
DOCX = os.path.join(BASE, "张哲铭TASK5.docx")
FONT = "宋体"
SIZE = Pt(10.5)

NAME = {
    "sh600900": "长江电力", "hk00700": "腾讯控股", "sh518880": "黄金ETF",
    "sh515450": "红利低波50ETF", "sz159941": "纳指ETF", "sh588000": "科创50ETF",
    "sh510300": "沪深300ETF", "sh510500": "中证500ETF",
}
CODES = list(NAME.keys())


def rfont(run, size=SIZE, bold=False, color=None, name=FONT):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = size
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def pfmt(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.alignment = align


def body(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False, indent=True, size=SIZE):
    p = doc.add_paragraph(); pfmt(p, align)
    if indent:
        p.paragraph_format.first_line_indent = Pt(21)
    rfont(p.add_run(text), size=size, bold=bold)
    return p


def title(doc, text, size=Pt(16)):
    p = doc.add_paragraph(); pfmt(p, WD_ALIGN_PARAGRAPH.CENTER)
    rfont(p.add_run(text), size=size, bold=True); return p


def heading(doc, text, size=Pt(13)):
    p = doc.add_paragraph(); pfmt(p, WD_ALIGN_PARAGRAPH.LEFT)
    rfont(p.add_run(text), size=size, bold=True); return p


def subheading(doc, text, size=Pt(11)):
    p = doc.add_paragraph(); pfmt(p, WD_ALIGN_PARAGRAPH.LEFT)
    p.paragraph_format.first_line_indent = Pt(21)
    rfont(p.add_run(text), size=size, bold=True); return p


def caption(doc, text):
    p = doc.add_paragraph(); pfmt(p, WD_ALIGN_PARAGRAPH.CENTER)
    rfont(p.add_run(text), size=Pt(9), bold=True); return p


def pic(doc, name, width=6.1):
    path = os.path.join(FIG, name)
    if not os.path.exists(path):
        body(doc, f"[缺图 {name}]"); return
    p = doc.add_paragraph(); pfmt(p, WD_ALIGN_PARAGRAPH.CENTER)
    p.add_run().add_picture(path, width=Inches(width))


def formula(doc, text):
    p = doc.add_paragraph(); pfmt(p, WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run(text)
    r.font.name = "Cambria Math"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(10.5); r.font.italic = True


def code_block(doc, code):
    for line in code.split("\n"):
        p = doc.add_paragraph(); pfmt(p, WD_ALIGN_PARAGRAPH.LEFT)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        r = p.add_run(line if line else " ")
        r.font.name = "Consolas"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def make_table(doc, headers, rows, size=Pt(8.5)):
    table = doc.add_table(rows=1, cols=len(headers)); table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        rfont(c.paragraphs[0].add_run(h), size=Pt(8.5), bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = cells[i].paragraphs[0].paragraph_format
            pf.space_before = Pt(0); pf.space_after = Pt(0)
            rfont(cells[i].paragraphs[0].add_run(str(v)), size=size)
    return table


def pct(x, d=1):
    try:
        return f"{float(x)*100:.{d}f}%"
    except Exception:
        return "-"


def build():
    clf = pd.read_csv(os.path.join(DATA, "clf_metrics.csv"))
    reg = pd.read_csv(os.path.join(DATA, "reg_metrics.csv"))
    cancer = pd.read_csv(os.path.join(DATA, "cancer_metrics.csv"))
    eda = pd.read_csv(os.path.join(DATA, "eda_summary.csv"))

    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    st.font.size = SIZE

    # ===== 封面 =====
    title(doc, "AI 交易引擎：机器学习算法与场景应用")
    info = doc.add_paragraph(); pfmt(info, WD_ALIGN_PARAGRAPH.CENTER)
    rfont(info.add_run("北京大学 AI 量化工作坊 · TASK5　　姓名：张哲铭"), size=Pt(10.5))
    doc.add_paragraph()

    body(doc,
        "本任务系统学习分类型机器学习算法及其在量化交易中的应用。全文首先讲清逻辑回归、决策树、随机森林、"
        "KNN 等基础分类算法，以及支持向量机、梯度提升树（GBDT/XGBoost/LightGBM）等在检索到的重要文献中被"
        "反复验证有效的算法，逐一给出原理与优缺点；随后解释混淆矩阵、准确率/精确率/召回率/F1、ROC 曲线与 AUC "
        "等模型评价指标；接着检索并综述机器学习在量化交易领域的重要学术与行业成果；然后进入 Python 实证：以 8 "
        "个代表性标的（长江电力、腾讯控股、黄金 ETF、红利低波 50ETF、纳指 ETF、科创 50ETF、沪深 300ETF、中证 "
        "500ETF）的近八年前复权日线为样本，讲清训练/验证/测试的时间序列划分、无未来函数的特征工程，构建并用"
        "网格搜索调优分类与回归两套模型，输出准确率/精确率/召回率/F1、混淆矩阵、ROC 曲线、AUC 以及 RMSE/MAE/"
        "R² 等完整指标对比，并配合乳腺癌数据集作为高信噪比对照。所有代码规避未来函数、绝不打乱时间顺序，"
        "为 TASK6 的策略构建奠定基础。")

    # ===== 一、分类型机器学习算法 =====
    heading(doc, "一、分类型机器学习算法解释")
    body(doc,
        "分类型机器学习算法的目标，是根据一组输入特征 X（如技术指标、财务因子）预测一个离散的类别标签 y（如"
        "“下一交易日上涨/不涨”）。下面按“从简单到复杂、从线性到集成”的顺序，介绍本任务实际使用的核心算法。")

    subheading(doc, "1. 逻辑回归（Logistic Regression）")
    body(doc,
        "逻辑回归是最基础的分类算法。它先对特征做线性加权 z = w·x + b，再用 Sigmoid 函数 σ(z)=1/(1+e^(−z)) 把"
        "结果压缩到 (0,1) 区间作为“属于正类的概率”，以 0.5 为默认阈值判类。它本质是广义线性模型，训练即最小化"
        "对数损失（交叉熵）。优点：模型简单、训练快、可解释性极强（每个特征的系数方向和大小直接反映影响）、"
        "天然输出概率，非常适合作为基线与后续策略的概率来源。缺点：只能刻画特征与对数几率之间的线性关系，"
        "无法自动捕捉非线性与特征交互，对共线性敏感、需要配合正则化（本文用 L2 正则 + 类别权重平衡）。")

    subheading(doc, "2. 决策树（Decision Tree）")
    body(doc,
        "决策树通过递归地对特征空间做二分切分来分类：每个内部节点选择“最能降低不纯度”（基尼系数或信息熵）的"
        "特征与阈值进行分裂，直到叶节点足够纯或达到深度限制。优点：可解释性强（可画出完整判断路径）、天然处理"
        "非线性与特征交互、对特征量纲不敏感、几乎不受多重共线性影响。缺点：单棵树极易过拟合、对训练样本扰动"
        "非常敏感（高方差），需要用剪枝、限制深度与叶节点最小样本数来控制——这也正是集成方法（随机森林、提升树）"
        "出现的动机。")

    subheading(doc, "3. 随机森林（Random Forest）")
    body(doc,
        "随机森林是决策树的 Bagging（自助采样聚合）集成：训练时对样本做有放回抽样、对特征做随机子集选择，"
        "并行训练大量彼此去相关的决策树，分类时按多数投票、回归时按均值汇总。通过“多样化 + 平均”大幅降低单树的"
        "方差。优点：精度高且稳健、抗过拟合、能给出特征重要性、对共线性和噪声不敏感、几乎无需特征标准化，是"
        "金融面板数据上最常用的“开箱即用”强模型。缺点：可解释性弱于单树、模型体积较大、对极高维稀疏数据不如"
        "线性模型。Gu、Kelly 与 Xiu（2020）在 3 万只美股上的大规模实证表明，树模型（随机森林、提升树）与神经"
        "网络显著优于线性方法，正是得益于其对非线性交互的刻画能力。")

    subheading(doc, "4. K 近邻（KNN, K-Nearest Neighbors）")
    body(doc,
        "KNN 是一种“惰性学习”算法：不显式训练模型，预测时找出与待判样本在特征空间中最近的 K 个训练样本，按它们"
        "的多数类别投票。优点：原理直观、无参数假设、能刻画任意复杂的局部决策边界。缺点：预测时需遍历全部训练"
        "样本、计算开销大；对特征尺度极其敏感（必须标准化）；在高维空间因“维度灾难”导致距离失效，且对类别不平衡"
        "和噪声敏感。本文将其作为对照算法，展示非参数方法在低信噪比金融数据上的表现。")

    subheading(doc, "5. 支持向量机（SVM）")
    body(doc,
        "支持向量机寻找一个能以最大间隔分开两类的超平面，并可通过核函数（本文用 RBF 核）把样本映射到高维空间以"
        "处理非线性可分问题。优点：在中小样本、高维、非线性问题上表现突出，泛化能力强、对局部噪声不敏感。缺点："
        "对大样本训练与概率估计计算开销大、核函数与惩罚参数 C 需要仔细调优、可解释性差。检索到的中文行业研究"
        "指出，SVM 在趋势预测与市场转折点识别中具有独特优势。")

    subheading(doc, "6. 梯度提升树（GBDT / XGBoost / LightGBM）")
    body(doc,
        "梯度提升是另一类树集成，但采用 Boosting 思想：串行地训练一系列弱决策树，每棵新树都去拟合前面所有树"
        "累计预测的“残差/负梯度”，从而不断纠错、逐步逼近目标。GBDT 是其经典实现；XGBoost 在其基础上加入二阶"
        "梯度、正则项与工程优化，训练更快更稳、抗过拟合；LightGBM 用直方图算法与叶子优先生长进一步提速，尤其"
        "适合大规模数据。优点：预测精度通常为传统方法之最、能自动捕捉复杂交互、内置正则、有特征重要性。缺点："
        "串行训练、超参数多、调优成本高、比随机森林更易过拟合噪声。Krauss、Do 与 Huck（2017）在标普 500 统计"
        "套利研究中发现，梯度提升树与随机森林、深度神经网络的集成能获得显著超额收益。AdaBoost 是最早的 Boosting "
        "算法，通过动态提高被错分样本的权重来串行改进，本文亦纳入对比。")

    # ===== 二、模型评价指标 =====
    heading(doc, "二、机器学习模型评价指标解释")
    subheading(doc, "1. 混淆矩阵（Confusion Matrix）")
    body(doc,
        "混淆矩阵是分类评估的基础，把预测结果按“真实类别 × 预测类别”交叉汇总为四格：真正例 TP（实际涨、预测涨）、"
        "真负例 TN（实际不涨、预测不涨）、假正例 FP（实际不涨、却预测涨，即“误报”）、假负例 FN（实际涨、却预测不涨，"
        "即“漏报”）。几乎所有分类指标都由这四个数派生而来。")
    make_table(doc, ["", "预测：涨(1)", "预测：不涨(0)"],
               [["实际：涨(1)", "TP 真正例", "FN 假负例(漏报)"],
                ["实际：不涨(0)", "FP 假正例(误报)", "TN 真负例"]])
    caption(doc, "表 2-1　二分类混淆矩阵结构")

    subheading(doc, "2. 准确率、精确率、召回率与 F1 分数")
    body(doc, "· 准确率 Accuracy = (TP+TN)/(TP+TN+FP+FN)，即整体判对的比例。类别不平衡时会失真（例如 73% 样本"
              "为上涨时，全猜涨也有 73% 准确率），故不能单看。", indent=True)
    formula(doc, "Accuracy = (TP + TN) / (TP + TN + FP + FN)")
    body(doc, "· 精确率 Precision = TP/(TP+FP)，即“预测为涨的样本里真正上涨的比例”，衡量信号的“纯度”，对应"
              "交易中的“出手命中率”，直接关系交易成本与胜率。", indent=True)
    formula(doc, "Precision = TP / (TP + FP)")
    body(doc, "· 召回率 Recall = TP/(TP+FN)，即“真正上涨的样本里被成功抓到的比例”，衡量“抓涨能力”，"
              "对应“机会捕获率”。", indent=True)
    formula(doc, "Recall = TP / (TP + FN)")
    body(doc, "· F1 分数 = 精确率与召回率的调和平均，兼顾两者，是类别不平衡下最常用的综合指标。", indent=True)
    formula(doc, "F1 = 2 × Precision × Recall / (Precision + Recall)")
    body(doc,
        "精确率与召回率通常此消彼长：把判正阈值调高，出手更谨慎、精确率升但召回降；调低则相反。交易场景往往"
        "更看重精确率（宁可少交易，也要提高每次出手的胜率），这一取舍在 TASK6 的“双阈值策略”中会被明确利用。")

    subheading(doc, "3. ROC 曲线与 AUC")
    body(doc,
        "ROC（受试者工作特征）曲线以“假正率 FPR = FP/(FP+TN)”为横轴、“真正率 TPR = 召回率”为纵轴，通过连续"
        "调整分类阈值（从 1 到 0）描出一条曲线，刻画模型在所有阈值下“抓对正例”与“误报负例”的权衡全貌。曲线越"
        "靠近左上角越好。AUC（曲线下面积）把这条曲线概括为一个 0~1 的数：其含义是“随机取一个正样本与一个负样本，"
        "模型给正样本打分更高的概率”。AUC=0.5 等于随机猜测、=1 为完美区分；金融日频涨跌预测因信噪比极低，AUC "
        "通常仅略高于 0.5，能稳定达到 0.53~0.58 已具备实用的边际信息。AUC 的最大优点是不依赖具体阈值、且对类别"
        "不平衡不敏感，因此是本文比较各分类模型的首选指标。")

    # ===== 三、文献综述 =====
    heading(doc, "三、重要文献与行业成果综述（量化交易中的机器学习算法）")
    body(doc,
        "为明确“在量化交易领域哪些机器学习算法真正重要、为何重要”，本文检索并梳理了代表性学术论文与行业成果，"
        "结论高度一致地指向“树集成与神经网络优于线性方法，而基础算法仍是不可或缺的基线与解释工具”。")
    body(doc,
        "（1）Gu、Kelly 与 Xiu（2020）《Empirical Asset Pricing via Machine Learning》（发表于顶级期刊 Review of "
        "Financial Studies，33(5):2223-2273）是该领域被引用最多的奠基性论文之一。作者在 1957–2016 年约 3 万只"
        "美股、近 900 个预测变量上，对 OLS、弹性网、随机森林、梯度提升树、神经网络等做“赛马式”对比，发现："
        "树模型与神经网络的样本外预测能力显著优于线性模型，其优势来源正是对“非线性与特征交互”的刻画；且所有"
        "方法一致认同的主导预测信号是动量、流动性与波动率。基于神经网络预测的多空组合样本外夏普比率超过 1.8，"
        "是最佳线性模型的两倍以上。这为本文“以量价动量/波动/流动性类因子为主、以树模型为主力”的设计提供了依据。", indent=True)
    body(doc,
        "（2）Krauss、Do 与 Huck（2017）《Deep neural networks, gradient-boosted trees, random forests: Statistical "
        "arbitrage on the S&P 500》（European Journal of Operational Research）在标普 500 成分股上系统比较了深度"
        "神经网络、梯度提升树与随机森林，并发现三者的“集成”在扣除成本后仍能取得显著的统计套利收益，凸显了树"
        "集成方法在真实交易中的实用价值。", indent=True)
    body(doc,
        "（3）综述类成果方面，Forecasting Stock Market Prices Using Machine Learning and Deep Learning Models"
        "（IJFS, 2023, 11(3):94）等系统综述总结出一条标准 ML 流水线：加载数据→预处理与特征选择→划分训练/测试→"
        "训练→评估→超参数调优；并明确“回归任务用 RMSE/MAE/R² 评估，分类任务用 Accuracy/Precision/Recall/F1，"
        "二分类另报 ROC/AUC”，最常用的算法为随机森林、XGBoost、SVM、LSTM，且集成/混合模型表现最佳。本文的"
        "评估体系与之完全对齐。", indent=True)
    body(doc,
        "（4）中文行业成果（如聚宽社区《机器学习算法在量化交易中的应用研究》等）指出：线性/逻辑回归是最早、最"
        "易解释的基线；SVM 擅长小样本、非线性与高维模式识别，利于捕捉市场转折；随机森林在处理高维金融数据、"
        "特征交互上表现最佳；XGBoost/LightGBM 等梯度提升在因子挖掘与选股中应用最广。这与英文文献的结论互为印证。", indent=True)
    body(doc,
        "综上，本任务选用的算法谱系——逻辑回归、决策树、随机森林、KNN（作业要求的基础算法）+ SVM、GBDT、"
        "AdaBoost、XGBoost、LightGBM（文献验证的重要算法）——既覆盖教学要求，又贴合行业前沿实践。")

    # ===== 四、数据与特征工程 =====
    heading(doc, "四、数据、特征工程与时间序列划分")
    subheading(doc, "1. 数据来源与研究标的")
    body(doc,
        "沿用本工作坊统一数据：经腾讯自选股接口获取的 8 个标的前复权日线（开高低收量），样本区间约 2018 年"
        "至 2026 年、多数标的近 2000 个交易日。标的覆盖个股与宽基/主题 ETF，风格互补：长江电力（低波红利个股）、"
        "腾讯控股（高波成长个股/港股）、黄金 ETF（避险商品）、红利低波 50ETF（低波防御）、纳指 ETF（海外成长）、"
        "科创 50ETF（高波科技）、沪深 300ETF（大盘）、中证 500ETF（中盘）。此外，对两只个股（长江电力、腾讯）"
        "额外补充了季度财务因子（EPS、营收/净利 TTM 或同比、ROE、净利率等），并严格按财务信息的“发布日”前向"
        "对齐到日频，杜绝“用尚未披露的财报预测过去”的未来函数；ETF 无对应财务报表，故仅用量价技术因子。")

    subheading(doc, "2. 特征工程（保证相关性、无未来函数、稳定性、可计算性）")
    body(doc,
        "共构造技术因子约 38 个（个股叠加财务因子后约 43 个），分六大类：①动量类（1/3/5/10/20/60 日累计收益、"
        "对数收益）；②均线偏离类（5/10/20/60 日乖离率、短长均线比值）；③波动率类（多窗口收益标准差、归一化 ATR、"
        "当日振幅）；④量价类（成交量变化、5/20 日量比、20 日量价相关性）；⑤经典技术指标（RSI6/14、MACD 的"
        " DIF/DEA/柱、KDJ 的 K/D/J、布林带 %B 与带宽）；⑥价格位置类（20/60 日价格分位、隔夜/日内收益）。"
        "工程上严守四条纪律：")
    body(doc, "· 相关性：优先选用与短期收益有经济学关联的因子（动量、波动、量价背离），并在 EDA 中量化每个"
              "因子与目标的相关性；", indent=True)
    body(doc, "· 无未来函数：所有滚动窗口只用当日及历史数据；标签用“下一交易日收益”定义在未来，特征与标签"
              "在时间上严格错位，构造后统一 dropna；财务因子按发布日对齐；", indent=True)
    body(doc, "· 稳定性：全部采用比率、对数收益、标准化乖离等无量纲/弱趋势特征，避免绝对价格随时间漂移导致的"
              "分布不稳定；", indent=True)
    body(doc, "· 可计算性：所有因子仅由 OHLCV（+个股财报）即可复现，无外部不可得数据。", indent=True)

    subheading(doc, "3. 预测目标（标签）定义")
    body(doc,
        "同时构造两套目标：分类目标 y_cls = “下一交易日收益 ≥ 0 记为 1（不跌），否则 0”；回归目标 y_reg = "
        "“下一交易日收益率”。之所以分类用“≥0”而非严格“>0”，是因为纳指、红利低波等跨境/低波 ETF 存在大量“平盘日”"
        "（收益恰为 0，占比可达 40%），若把平盘归为“跌”会造成标签严重失衡且不符合“不跌即可持有”的择时语义。")

    subheading(doc, "4. 训练/验证/测试划分（时间序列，绝不打乱）")
    body(doc,
        "金融数据是时间序列，绝不能像普通静态数据那样随机打乱划分，否则会用“未来样本”训练、再回头预测“过去”，"
        "造成严重的信息泄漏和虚高业绩。本文严格按时间顺序把每个标的切成 60% 训练集 / 20% 验证集 / 20% 测试集，"
        "保证训练集时间完全早于验证集、验证集完全早于测试集；网格搜索调参时进一步在“训练+验证”内部使用"
        " TimeSeriesSplit（前向扩窗交叉验证），每一折都用较早的数据训练、较晚的数据验证；标准化所用的均值方差"
        "只在训练集上估计，再套用到验证/测试集。测试集在整个建模过程中完全隔离，只在最后评估时使用一次。")
    # 划分表
    rows = []
    for _, r in eda.iterrows():
        n = int(r["n"])
        rows.append([NAME.get(r["code"], r["code"]), n,
                     int(n*0.6), int(n*0.2), n-int(n*0.6)-int(n*0.2),
                     int(r["n_feat"]), pct(r["up_ratio"])])
    make_table(doc, ["标的", "有效样本", "训练", "验证", "测试", "特征数", "上涨占比"], rows)
    caption(doc, "表 4-1　各标的样本量、时间序列划分与类别分布")

    # ===== 五、EDA =====
    heading(doc, "五、探索性数据分析与特征诊断")
    body(doc,
        "在建模前，对每个标的做完整的探索性数据分析（EDA），以理解目标分布、特征与目标的关系、特征间的冗余"
        "（多重共线性）等。受篇幅所限，正文以长江电力与纳指 ETF 为例展示，其余标的图表见附带的 figures 目录与"
        "网页版报告。")

    subheading(doc, "1. 目标变量分布")
    pic(doc, "eda_target_sh600900.png", 6.2)
    caption(doc, "图 5-1　长江电力：下期涨跌次数分布（左，分类目标）与下期收益率分布（右，回归目标）")
    body(doc,
        "长江电力下期涨跌大致均衡（上涨占比约 50.6%），收益率分布近似以 0 为中心、尖峰厚尾的钟形，符合日频收益"
        "的典型特征。作为对照，纳指 ETF 因大量平盘日，分类目标偏向“不跌”一侧，收益率分布更集中。")
    pic(doc, "eda_target_sz159941.png", 6.2)
    caption(doc, "图 5-2　纳指 ETF：下期涨跌次数分布与下期收益率分布")

    subheading(doc, "2. 特征与目标的相关性及排序")
    pic(doc, "eda_topcorr_sh600900.png", 5.8)
    caption(doc, "图 5-3　长江电力：与下期收益率相关性最高的 Top15 特征（红正绿负，按|相关系数|排序）")
    body(doc,
        "把所有特征对目标（下期收益率）的皮尔逊相关系数按绝对值从大到小排序可见：日内收益（intraday）、"
        "隔夜收益（overnight）、当日振幅（hl_range）、短期量比等“反转/量价”类因子的相关性最高，但绝对值普遍"
        "低于 0.2。这正是股票日频收益“低信噪比”的真实写照——单个特征的线性预测力很弱，必须依靠机器学习模型"
        "对众多弱信号做非线性组合。各标的完整的相关性排序（X 对 Y 的相关系数）已导出为 data/corr_rank_*.csv。")

    subheading(doc, "3. 特征相关性矩阵与多重共线性")
    pic(doc, "eda_corrmat_sh600900.png", 5.6)
    caption(doc, "图 5-4　长江电力：Top20 相关特征的相关性矩阵热力图")
    body(doc,
        "相关性矩阵显示，特征间存在明显的“成组高相关”：KDJ 的 K/D/J 三线彼此相关系数常超过 0.9，MACD 的 "
        "DIF/DEA/柱、不同窗口的均线乖离与动量之间也高度共线。我们用两种方式量化多重共线性：一是方差膨胀因子"
        " VIF，二是统计 |相关系数|>0.9 的高相关特征对。结果显示每个标的都存在 15~22 对强共线特征、部分 VIF 极高"
        "（KDJ/均线类）。处理方式为：①主力模型选用树集成（随机森林、GBDT/XGBoost），它们基于特征分裂、对共线性"
        "天然不敏感，无需剔除；②对线性/逻辑回归施加 L2 正则以稳定系数；③在解释特征重要性时，对成组共线特征"
        "合并解读，避免把同一信息重复计数。相关诊断结果见 data/vif_*.csv 与 data/highcorr_pairs_*.csv。")

    subheading(doc, "4. 关键特征按目标分组的箱线图")
    pic(doc, "eda_box_sh600900.png", 6.2)
    caption(doc, "图 5-5　长江电力：Top6 关键特征在“下期涨”与“下期跌/平”两组下的箱线图对比")
    body(doc,
        "箱线图对比两类样本下关键特征的分布差异：若某特征在“涨”与“跌”两组间中位数和箱体错开明显，说明该特征"
        "具备判别力。可以看到日内收益、振幅等特征在两组间确有可见差异（体现短期反转效应），但重叠区域依然很大，"
        "再次印证单特征区分度有限、需要多特征联合建模。")

    # ===== 六、建模与评估 =====
    heading(doc, "六、模型构建、网格调优与评估")
    subheading(doc, "1. 分类模型评估指标对比")
    body(doc,
        "对每个标的、每个分类算法，在“训练+验证”集上用 TimeSeriesSplit + 网格搜索择优，再在独立测试集上评估"
        "准确率、精确率、召回率、F1 与 AUC，并输出混淆矩阵与 ROC 曲线。下表汇总各算法在 8 个标的上的平均表现"
        "（按平均 AUC 降序）。")
    g = clf.groupby("model")[["accuracy", "precision", "recall", "f1", "auc"]].mean().sort_values("auc", ascending=False)
    rows = [[m, pct(r["accuracy"]), pct(r["precision"]), pct(r["recall"]),
             pct(r["f1"]), f"{r['auc']:.3f}"] for m, r in g.iterrows()]
    make_table(doc, ["分类算法", "准确率", "精确率", "召回率", "F1", "AUC(均)"], rows)
    caption(doc, "表 6-1　各分类算法在 8 标的上的平均评估指标（按 AUC 降序）")
    best_clf_model = g.index[0]
    best_clf_auc = g.iloc[0]["auc"]
    body(doc,
        f"整体来看，各模型平均 AUC 集中在 0.5 附近略偏上，其中表现相对最好的是“{best_clf_model}”（平均 AUC ≈ "
        f"{best_clf_auc:.3f}）。这一结果诚实地反映了金融日频方向预测的极高难度：即便是集成模型，样本外 AUC 也仅"
        "略优于随机猜测。但正如 Gu-Kelly-Xiu 所指出的，哪怕是很小的稳定edge，经过合理的仓位与风控放大，仍可能"
        "转化为有意义的策略收益——这正是 TASK6 要检验的。")
    pic(doc, "model_compare.png", 6.4)
    caption(doc, "图 6-1　分类模型平均 AUC/F1 对比（左）与回归模型平均方向命中率对比（右）")

    subheading(doc, "2. ROC 曲线与混淆矩阵（以长江电力为例）")
    pic(doc, "roc_sh600900.png", 5.2)
    caption(doc, "图 6-2　长江电力：各分类模型 ROC 曲线对比")
    pic(doc, "cm_sh600900.png", 4.4)
    caption(doc, "图 6-3　长江电力：最优分类模型的混淆矩阵")
    body(doc,
        "ROC 曲线越靠近左上、AUC 越大越好；对角虚线为随机基准。混淆矩阵则直观展示模型在测试集上的四类判别数量，"
        "可据此计算精确率/召回率。各标的的 ROC 与混淆矩阵图完整存于 figures 目录。")

    subheading(doc, "3. 回归模型评估指标对比")
    body(doc,
        "回归任务直接预测“下期收益率”，用 RMSE（均方根误差）、MAE（平均绝对误差）、R²（决定系数）评估拟合优度，"
        "并额外报告“方向命中率”（预测涨跌方向与真实一致的比例），因为对交易而言方向比数值更重要。")
    gr = reg.groupby("model")[["rmse", "mae", "r2", "dir_acc"]].mean().sort_values("dir_acc", ascending=False)
    rows = [[m, f"{r['rmse']:.4f}", f"{r['mae']:.4f}", f"{r['r2']:.4f}", pct(r["dir_acc"])]
            for m, r in gr.iterrows()]
    make_table(doc, ["回归算法", "RMSE(均)", "MAE(均)", "R²(均)", "方向命中(均)"], rows)
    caption(doc, "表 6-2　各回归算法在 8 标的上的平均评估指标（按方向命中率降序）")
    body(doc,
        "回归结果同样体现低信噪比特征：各模型 R² 普遍接近 0 甚至微负（说明对收益率数值的解释力极弱），但方向"
        "命中率仍能稳定在 50% 附近略高。这说明“预测具体涨幅”几乎不可能，而“预测涨跌方向”尚存一线微弱但可利用"
        "的信息——这也是 TASK6 采用“分类概率 + 阈值”而非“回归数值”来驱动策略的根本原因。")

    subheading(doc, "4. 逐标的最优模型一览")
    rows = []
    for code in CODES:
        sub = clf[clf["code"] == code]
        if len(sub) == 0:
            continue
        b = sub.sort_values("auc", ascending=False).iloc[0]
        rows.append([NAME[code], b["model"], f"{b['auc']:.3f}",
                     pct(b["accuracy"]), pct(b["precision"]),
                     pct(b["recall"]), pct(b["f1"])])
    make_table(doc, ["标的", "最优分类模型", "AUC", "准确率", "精确率", "召回率", "F1"], rows)
    caption(doc, "表 6-3　各标的表现最优的分类模型（按测试集 AUC）")

    # ===== 七、乳腺癌对照 =====
    heading(doc, "七、乳腺癌数据集分类示例（高信噪比对照）")
    body(doc,
        "为直观说明“同样的算法，数据可预测性决定模型上限”，本文用 scikit-learn 自带的威斯康星乳腺癌数据集"
        "（569 样本、30 个细胞核形态特征、二分类：恶性/良性）做标准分类流水线。该数据特征与标签因果关系明确、"
        "信噪比高，四种基础算法均取得优异表现：")
    rows = [[r["model"], pct(r["accuracy"]), pct(r["precision"]), pct(r["recall"]),
             pct(r["f1"]), f"{r['auc']:.3f}"] for _, r in cancer.iterrows()]
    make_table(doc, ["算法", "准确率", "精确率", "召回率", "F1", "AUC"], rows)
    caption(doc, "表 7-1　乳腺癌数据集上各分类算法评估指标")
    pic(doc, "cancer_demo.png", 6.4)
    caption(doc, "图 7-1　乳腺癌数据集：随机森林混淆矩阵（左）与各模型 ROC 曲线（右）")
    body(doc,
        "对比可见：乳腺癌数据上 AUC 普遍高达 0.94~0.998、准确率超过 94%；而同样的算法用在股票日频涨跌上 AUC 仅"
        "略高于 0.5。这一强烈反差生动说明——机器学习不是“点石成金”，模型效果的天花板首先由数据本身的可预测性"
        "决定。金融市场的高效率与强噪声，使得任何模型都只能提取微弱的统计优势，这也要求我们在策略层面用严格的"
        "风控与仓位管理来把“微弱edge”转化为“稳健收益”。")

    # ===== 八、结论 =====
    heading(doc, "八、结论与对 TASK6 的衔接")
    body(doc, "1. 算法层面：树集成方法（随机森林、GBDT/XGBoost/LightGBM）在 8 个标的上综合表现最稳，"
              "与 Gu-Kelly-Xiu(2020)、Krauss et al.(2017) 等文献结论一致；逻辑回归/决策树/KNN 作为基线，"
              "可解释性强但预测力有限；SVM 计算成本高、优势不明显。", indent=True)
    body(doc, "2. 数据层面：股票日频方向预测信噪比极低（AUC 约 0.5~0.58、回归 R²≈0），单特征线性相关性普遍<0.2，"
              "且存在严重多重共线性——这要求以树模型为主、以 AUC/F1 而非准确率为主评估指标。", indent=True)
    body(doc, "3. 工程层面：严格的时间序列划分、无未来函数的特征工程、财务因子按发布日对齐，是结果可信的前提；"
              "乳腺癌对照进一步说明“数据可预测性决定模型上限”。", indent=True)
    body(doc, "4. 衔接 TASK6：既然“预测涨跌方向”比“预测涨幅”更可行，TASK6 将采用分类模型输出的“上涨概率”，"
              "配合双阈值、概率加权仓位、技术指标过滤与止损止盈，把本任务得到的微弱预测edge转化为可回测的交易"
              "策略，并系统对比不同算法、不同参数的策略效果。", indent=True)

    body(doc, "（注：本文为量化学习实践，不构成任何投资建议；市场有风险，决策需谨慎。）", indent=True)

    doc.save(DOCX)
    print(f"已生成 {DOCX}")
    return DOCX


if __name__ == "__main__":
    build()
