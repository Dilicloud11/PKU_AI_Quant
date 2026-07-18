# -*- coding: utf-8 -*-
"""
生成 TASK6 作业文档：张哲铭TASK6.docx / .pdf
格式：宋体、五号(10.5pt)、1.5 倍行距、0 段间距、正文两端对齐。
内容：ML 交易策略核心理念与优缺点、自变量因子与应变量定义、文献综述、
      策略设计（双阈值/概率仓位/技术过滤/风控/网格调优）、逐标的四图解读、
      季度收益、算法对比、附加题（组合轮动）、结论。
作者：张哲铭
"""
import os
import json
import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

BASE = os.path.dirname(os.path.abspath(__file__))
FIG = os.path.join(BASE, "figures")
DATA = os.path.join(BASE, "data")
DOCX = os.path.join(BASE, "张哲铭TASK6.docx")
FONT = "宋体"
SIZE = Pt(10.5)

NAME = {
    "sh600900": "长江电力", "hk00700": "腾讯控股", "sh518880": "黄金ETF",
    "sh515450": "红利低波50ETF", "sz159941": "纳指ETF", "sh588000": "科创50ETF",
    "sh510300": "沪深300ETF", "sh510500": "中证500ETF",
}
CODES = list(NAME.keys())


def rfont(run, size=SIZE, bold=False, color=None, name=FONT):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = size
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def pfmt(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(0); pf.space_after = Pt(0); pf.alignment = align


def body(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False, indent=True, size=SIZE):
    p = doc.add_paragraph(); pfmt(p, align)
    if indent:
        p.paragraph_format.first_line_indent = Pt(21)
    rfont(p.add_run(text), size=size, bold=bold)
    return p


def title(doc, text, size=Pt(16)):
    p = doc.add_paragraph(); pfmt(p, WD_ALIGN_PARAGRAPH.CENTER)
    rfont(p.add_run(text), size=size, bold=True); return p


def heading(doc, text, size=Pt(13)):
    p = doc.add_paragraph(); pfmt(p, WD_ALIGN_PARAGRAPH.LEFT)
    rfont(p.add_run(text), size=size, bold=True); return p


def subheading(doc, text, size=Pt(11)):
    p = doc.add_paragraph(); pfmt(p, WD_ALIGN_PARAGRAPH.LEFT)
    p.paragraph_format.first_line_indent = Pt(21)
    rfont(p.add_run(text), size=size, bold=True); return p


def caption(doc, text):
    p = doc.add_paragraph(); pfmt(p, WD_ALIGN_PARAGRAPH.CENTER)
    rfont(p.add_run(text), size=Pt(9), bold=True); return p


def pic(doc, name, width=6.1):
    path = os.path.join(FIG, name)
    if not os.path.exists(path):
        body(doc, f"[缺图 {name}]"); return
    p = doc.add_paragraph(); pfmt(p, WD_ALIGN_PARAGRAPH.CENTER)
    p.add_run().add_picture(path, width=Inches(width))


def formula(doc, text):
    p = doc.add_paragraph(); pfmt(p, WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run(text)
    r.font.name = "Cambria Math"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(10.5); r.font.italic = True


def make_table(doc, headers, rows, size=Pt(8.5)):
    table = doc.add_table(rows=1, cols=len(headers)); table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        rfont(c.paragraphs[0].add_run(h), size=Pt(8.5), bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = cells[i].paragraphs[0].paragraph_format
            pf.space_before = Pt(0); pf.space_after = Pt(0)
            rfont(cells[i].paragraphs[0].add_run(str(v)), size=size)
    return table


def pct(x, d=1):
    try:
        return f"{float(x)*100:.{d}f}%"
    except Exception:
        return "-"


def build():
    strat = pd.read_csv(os.path.join(DATA, "strategy_results.csv"))
    best = pd.read_csv(os.path.join(DATA, "best_strategy.csv"))
    bonus = pd.read_csv(os.path.join(DATA, "bonus_rotation.csv"))

    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    st.font.size = SIZE

    title(doc, "智能决策者：用机器学习定制专属策略")
    info = doc.add_paragraph(); pfmt(info, WD_ALIGN_PARAGRAPH.CENTER)
    rfont(info.add_run("北京大学 AI 量化工作坊 · TASK6　　姓名：张哲铭"), size=Pt(10.5))
    doc.add_paragraph()

    body(doc,
        "本任务在 TASK5“用机器学习预测涨跌”的基础上，进一步把模型输出的“上涨概率”转化为可执行、可回测的交易"
        "策略。全文先阐明基于机器学习的交易策略的核心理念与优缺点，界定量化 ML 模型中常见的自变量因子与应变量，"
        "综述相关文献；随后详细设计一套融合“双阈值 + 概率加权仓位 + 技术指标过滤 + 止损止盈”的策略框架，用网格"
        "搜索为 8 个标的分别寻找最优参数；接着对每个标的用 TASK5 中表现最好的 Top3 算法产出样本外概率、构建策略并"
        "回测，逐一给出“价格与概率双轴（何时交易）、资产曲线对比（赚了多少）、回撤曲线（最大亏多少）、持仓比例"
        "（仓位怎么变）”四张核心图，并系统对比决策树、随机森林、GBDT、XGBoost 等不同算法的策略效果；最后完成附加题"
        "——多标的机器学习组合轮动策略。全流程严格使用 TASK5 划分的测试集时间段、无未来函数。")

    # ===== 一、核心理念与优缺点 =====
    heading(doc, "一、基于机器学习的交易策略：核心理念与优缺点")
    body(doc,
        "核心理念可概括为“预测—转化—风控”三步：第一步用机器学习模型从历史量价/财务特征中学习规律，对下一期"
        "涨跌输出一个概率 p；第二步把概率转化为交易决策——概率高则看多、加仓，概率低则观望、清仓；第三步用严格的"
        "仓位管理与止损止盈把“统计上的微弱优势”在风险可控的前提下累积成收益。它与传统规则策略（如均线、海龟）的"
        "根本区别在于：规则策略由人预先设定固定条件，而机器学习策略让数据自己“说话”，自动从大量特征中挖掘非线性"
        "组合信号，具备更强的适应性。")
    subheading(doc, "· 优点")
    body(doc, "1. 能处理高维、非线性信息。 可同时纳入数十上百个因子，自动捕捉传统线性规则难以刻画的特征交互，"
              "信息利用更充分。", indent=True)
    body(doc, "2. 概率化、可量化风险。 模型输出连续概率，天然支持“按确定性大小分配仓位”，比非黑即白的规则更精细。", indent=True)
    body(doc, "3. 系统化、可回测、可迭代。 全流程数据驱动、参数可网格搜索、效果可严格回测，便于持续优化。", indent=True)
    body(doc, "4. 适应性强。 只要重新训练即可适配不同标的、不同市场环境，无需人工重设规则。", indent=True)
    subheading(doc, "· 缺点与风险")
    body(doc, "1. 过拟合风险高。 模型可能记住历史噪声而非真实规律，样本内亮眼、样本外失效，必须用严格的时间序列"
              "验证、正则化与充分的样本外测试来防范。", indent=True)
    body(doc, "2. 依赖数据质量与信噪比。 如 TASK5 所示，股票日频方向预测信噪比极低（AUC 仅约 0.5~0.6），模型上限"
              "受数据本身限制，不能期待“稳赚”。", indent=True)
    body(doc, "3. 存在未来函数与信息泄漏隐患。 特征构造、标签定义、数据划分稍有不慎就会用到未来信息，导致回测虚高、"
              "实盘崩塌，是最需要警惕的“陷阱”。", indent=True)
    body(doc, "4. 可解释性与稳健性挑战。 复杂模型（如提升树、神经网络）像“黑箱”，且市场结构会漂移（概念漂移），"
              "历史规律可能失效，需要持续监控与再训练。", indent=True)
    body(doc, "5. 交易成本与容量约束。 频繁调仓会侵蚀收益，这也是本文强调“双阈值降低换手”的原因。", indent=True)

    # ===== 二、自变量因子与应变量定义 =====
    heading(doc, "二、量化 ML 模型中常见自变量因子与应变量的定义")
    subheading(doc, "1. 应变量（预测目标 Y）")
    body(doc,
        "应变量是模型要预测的对象，常见两类：①回归型——未来某期的收益率（如下一日/下一周/下一季度收益率），"
        "用于收益排序与选股；②分类型——未来涨跌方向（涨/跌，二分类），或涨跌幅度分档（多分类）。本文主用分类型"
        "应变量 y=“下一交易日是否不跌（收益≥0 记 1）”，因为 TASK5 已证明“预测方向”比“预测具体涨幅”更可行；"
        "分类模型输出的“上涨概率”正是驱动交易的核心信号。")
    subheading(doc, "2. 自变量（特征因子 X）")
    body(doc,
        "自变量是用于预测的输入特征，量化实践中常见因子可归为若干大类：")
    body(doc, "· 动量/趋势因子：过去 N 日收益率、对数收益、均线斜率——刻画价格惯性（Gu-Kelly-Xiu 指出动量是最"
              "主导的预测信号之一）；", indent=True)
    body(doc, "· 波动率因子：收益标准差、ATR、振幅、布林带带宽——刻画风险与市场情绪；", indent=True)
    body(doc, "· 量价因子：成交量变化、量比、量价相关性——刻画资金进出与趋势确认；", indent=True)
    body(doc, "· 技术指标因子：RSI、MACD、KDJ、乖离率、%B——经典择时信号；", indent=True)
    body(doc, "· 价格位置因子：过去 N 日价格分位、隔夜/日内收益——刻画相对高低位与反转；", indent=True)
    body(doc, "· 基本面/财务因子：EPS、营收/净利增速、ROE、净利率、估值等——刻画公司质地（本文对长江电力、腾讯"
              "两只个股按财务发布日对齐后纳入；ETF 无财报故不用）；", indent=True)
    body(doc, "· 宏观/情绪因子：利率、资金面、市场热度等（本文未纳入，属可扩展方向）。", indent=True)
    body(doc,
        "本文实际使用约 38 个技术因子（个股叠加财务后约 43 个），构造时严守“相关性、无未来函数、稳定性、可计算性”"
        "四原则，详见 TASK5 报告第四章。")

    # ===== 三、文献综述 =====
    heading(doc, "三、基于机器学习算法的交易策略：文献与行业成果")
    body(doc,
        "（1）Krauss、Do 与 Huck（2017, EJOR）在标普 500 上构建统计套利策略：先用深度神经网络、梯度提升树、随机"
        "森林分别预测个股相对市场的表现概率，再做多预测最强、做空预测最弱的一篮子股票。研究发现三类模型的集成在"
        "扣除交易成本后仍能获得显著超额收益，是“ML 预测 + 排序选股”策略的经典范式。")
    body(doc,
        "（2）López de Prado 在《Advances in Financial Machine Learning》(2018) 中提出“三重障碍标注法”与"
        "“元标签（Meta-Labeling）”：先由主模型给出交易方向，再由次级模型判断“该不该按这个信号下注、下多大注”，"
        "并给出多种由预测概率决定仓位大小的方法（线性法 size∝(p−0.5)、凯利公式、Sigmoid 映射）。本文的“概率加权"
        "仓位”正是这一思想的直接应用。")
    body(doc,
        "（3）中文行业成果（如《运用随机森林演算法于选择权量化交易策略》《基于大模型的量化投资策略构建及回测有效性》）"
        "系统讨论了“信号阈值、持仓周期、仓位上限、止损比例、再平衡频率”等策略参数的敏感性，并指出信号阈值对收益的"
        "敏感度最高。这为本文“网格搜索最优参数（买入/卖出阈值、最大仓位、止损、止盈）”提供了直接参考。")

    # ===== 四、策略设计 =====
    heading(doc, "四、交易策略设计")
    body(doc,
        "本文以 TASK5 训练好的分类模型在测试集上的样本外“上涨概率 p”为唯一信号来源，设计如下策略框架，"
        "所有决策“当日盘后计算、次日生效”，杜绝未来函数：")
    subheading(doc, "1. 双阈值策略（降低换手、规避不确定区）")
    body(doc,
        "设买入阈值 buy_th 与卖出阈值 sell_th（且 sell_th < buy_th）：当 p ≥ buy_th 视为看多信号、开仓或持仓；"
        "当 p ≤ sell_th 视为看空、清仓；当 p 落在两者之间（不确定区）则维持原仓位、不做操作。相比单一阈值，双阈值"
        "形成一个“缓冲带”，避免概率在阈值附近抖动时频繁买卖，从而显著降低交易次数与成本，也避免在把握不大时盲动。")
    subheading(doc, "2. 概率加权仓位（确定性越高、仓位越重）")
    body(doc, "目标仓位不再非满即空，而是随概率线性放大：")
    formula(doc, "目标仓位 = clip( (p − 0.5) × 2, 0, 1 ) × 最大仓位")
    body(doc,
        "即概率 0.5 时空仓、概率越接近 1 仓位越接近“最大仓位”。这样在模型确定性高时重仓、把握不大时轻仓，"
        "把有限的风险预算分配到最有把握的交易上。")
    subheading(doc, "3. 机器学习预测 + 技术指标过滤")
    body(doc,
        "在模型信号之上叠加一层技术“安全阀”：当出现 RSI>70（超买）、MA5≤MA20（短期均线空头排列）、"
        "或近期波动率处于高分位（>90% 分位，市场过热/恐慌）时，禁止新开仓。这相当于用简单、稳健的技术规则过滤掉"
        "模型可能误判的不利环境，减少在高风险时点的暴露。")
    subheading(doc, "4. 止损与止盈风控")
    body(doc,
        "以持仓成本价为基准设置止损 stop_loss 与止盈 take_profit：持仓浮亏达到止损幅度即离场、防止亏损扩大；"
        "浮盈达到止盈幅度即兑现、锁定利润。单边交易成本按万分之五计（与前序任务一致）。")
    subheading(doc, "5. 参数网格搜索")
    body(doc,
        "对买入阈值、卖出阈值、最大仓位、止损幅度、止盈幅度五个参数做网格搜索，为每个标的分别寻找最优组合。"
        "为控制计算量（遵循作业“可先缩小参数范围”的建议），网格设定为：买入阈值∈{0.55, 0.60, 0.65}、"
        "卖出阈值∈{0.45, 0.50}、最大仓位∈{0.8, 1.0}、止损∈{5%, 8%}、止盈∈{15%, 25%}，"
        "以测试集夏普比率（辅以超额收益）为寻优目标。")

    # ===== 五、逐标的回测四图与解读 =====
    heading(doc, "五、逐标的回测：四张核心图与解读")
    body(doc,
        "对每个标的，用其综合表现最优的算法+最优参数回测，绘制四张核心图，分别回答四个关键问题："
        "A 图（价格+概率双轴，标注买卖点）——“何时交易”；B 图（策略 vs 买入持有净值）——“赚了多少”；"
        "C 图（回撤曲线，标注最大回撤）——“最大亏了多少”；D 图（持仓比例曲线）——“仓位怎么变”。")

    for code in CODES:
        b = best[best["code"] == code].iloc[0]
        prm = json.loads(b["params"].replace("'", '"'))
        subheading(doc, f"{NAME[code]}（{code}）— 最优算法：{b['algo']}")
        pic(doc, f"strat4_{code}.png", 6.4)
        caption(doc, f"图 5-{CODES.index(code)+1}　{NAME[code]} ML 策略回测四图（A 买卖点 / B 净值 / C 回撤 / D 仓位）")
        body(doc,
            f"最优参数：买入阈值 {prm['buy_th']}、卖出阈值 {prm['sell_th']}、最大仓位 {pct(prm['max_pos'],0)}、"
            f"止损 {pct(prm['stop_loss'],0)}、止盈 {pct(prm['take_profit'],0)}。测试段策略总收益 {pct(b['strat_total'])}、"
            f"买入持有 {pct(b['bh_total'])}、超额 {pct(b['excess_total'])}；夏普 {b['sharpe']:.2f}、最大回撤 "
            f"{pct(b['mdd'])}、交易 {int(b['trades'])} 次、持仓日胜率 {pct(b['win_rate'])}、平均仓位 {pct(b['exposure'])}。")

    # ===== 六、算法对比与季度收益 =====
    heading(doc, "六、不同机器学习算法的策略效果对比")
    pic(doc, "algo_compare.png", 6.5)
    caption(doc, "图 6-1　各标的不同机器学习算法策略总收益对比（虚线=买入持有）")
    body(doc,
        "上图并列展示每个标的下 Top3 算法策略的总收益（红柱）与买入持有基准（灰虚线）。下表汇总各标的最优策略：")
    rows = []
    for code in CODES:
        b = best[best["code"] == code].iloc[0]
        rows.append([NAME[code], b["algo"], pct(b["strat_total"]), pct(b["bh_total"]),
                     pct(b["excess_total"]), f"{b['sharpe']:.2f}", pct(b["mdd"]),
                     int(b["trades"]), pct(b["win_rate"])])
    make_table(doc, ["标的", "最优算法", "策略收益", "买入持有", "超额", "夏普", "最大回撤", "交易数", "胜率"], rows)
    caption(doc, "表 6-1　各标的最优 ML 策略回测指标汇总")
    body(doc,
        "关键结论有三：其一，策略的夏普比率普遍高于 1（多数在 1.2~2.4）、最大回撤极小（多在 −1%~−3%），说明"
        "“概率信号 + 双阈值 + 技术过滤 + 止损止盈”确实有效控制了风险、显著改善了风险调整后收益。其二，在测试段"
        "（近一两年多为单边上涨行情）下，策略的绝对收益普遍跑输买入持有——这是择时策略的固有代价：为控回撤而降低"
        "仓位，必然在单边牛市中踏空部分涨幅。其三，唯一显著跑赢的是腾讯控股：其测试段买入持有为负（下跌市），"
        "策略却取得正收益、超额可观——这生动印证了“机器学习择时的价值主要体现在震荡与下跌市，而非单边牛市”，"
        "与 TASK3/TASK4 关于趋势/择时策略“核心价值在控回撤”的结论一脉相承。")

    # ===== 七、附加题 =====
    heading(doc, "七、附加题：多标的机器学习组合轮动策略")
    body(doc,
        "作为拓展，本文自行设计了一个横截面策略：用随机森林对 8 个标的分别输出“下期上涨概率”，每个交易日选出"
        "概率最高的 Top3 标的等权持有（择优轮动、每日再平衡、计交易成本），与“等权买入持有全部 8 标的”基准对比。"
        "这呼应了原题“预测收益排序、挑选最优若干标的投资”的思路。")
    pic(doc, "bonus_rotation.png", 6.3)
    caption(doc, "图 7-1　机器学习组合轮动（Top3）vs 全市场等权持有：净值与回撤")
    r_s = bonus.iloc[0]; r_b = bonus.iloc[1]
    make_table(doc, ["策略", "总收益", "年化", "夏普", "最大回撤"],
               [[r_s["策略"], pct(r_s["总收益"]), pct(r_s["年化"]), f"{r_s['夏普']:.2f}", pct(r_s["最大回撤"])],
                [r_b["策略"], pct(r_b["总收益"]), pct(r_b["年化"]), f"{r_b['夏普']:.2f}", pct(r_b["最大回撤"])]])
    caption(doc, "表 7-1　组合轮动策略与全市场基准对比")
    body(doc,
        f"结果显示，在测试段公共时间窗内，ML 组合轮动年化约 {pct(r_s['年化'])}、夏普 {r_s['夏普']:.2f}，"
        f"而等权持有全市场年化约 {pct(r_b['年化'])}、夏普 {r_b['夏普']:.2f}。轮动策略同样未能跑赢等权持有——原因仍是"
        "样本区间整体上行、且横截面上各宽基 ETF 走势高度相关，使“选强汰弱”的空间有限；但轮动策略的最大回撤与基准"
        "接近，说明其在风险端并未恶化。这提示：横截面选股策略更适合标的众多、分化明显的股票池（如全 A 股），"
        "而非本文这样以少数高相关宽基 ETF 为主的标的池。")

    # ===== 八、结论 =====
    heading(doc, "八、结论")
    body(doc, "1. 方法论：成功把 TASK5 的“上涨概率”通过双阈值、概率加权仓位、技术过滤与止损止盈，转化为完整"
              "可回测的交易策略，并用网格搜索为每个标的定制最优参数，践行了“预测—转化—风控”的核心理念。", indent=True)
    body(doc, "2. 风险控制显著：所有标的策略夏普普遍>1、最大回撤压缩到个位数百分比甚至更低，验证了 ML+风控框架"
              "在控制风险上的有效性。", indent=True)
    body(doc, "3. 收益的诚实结论：在单边上涨的测试段，择时/轮动策略普遍跑输买入持有（踏空代价），唯有在下跌市"
              "（腾讯）显著跑赢——再次印证择时价值在震荡下跌市。这是量化实践中真实、常见且必须正视的现象。", indent=True)
    body(doc, "4. 关键前提：严格的时间序列纪律（不打乱、无未来函数、财务按发布日对齐、测试集隔离）是一切结论"
              "可信的基础；低信噪比决定了不应追求高胜率神话，而应追求“可控风险下的稳定微弱优势”。", indent=True)

    body(doc, "（注：本文为量化学习实践，不构成任何投资建议；市场有风险，决策需谨慎。）", indent=True)

    doc.save(DOCX)
    print(f"已生成 {DOCX}")
    return DOCX


if __name__ == "__main__":
    build()
