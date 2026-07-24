# -*- coding: utf-8 -*-
"""
生成 TASK8 综合学习报告：张哲铭TASK8.docx / .pdf
格式：宋体、五号(10.5pt)、1.5 倍行距、0 段间距、正文两端对齐。
结构：封面 / 目录 / 摘要 / 正文(量化核心概念·策略综合分析·机器学习应用总结·结论展望) / 附录(改进建议)。
图表全文统一编号（图1、图2…；表1、表2…），附录建议逐项编号，正文按编号引用。
作者：张哲铭
"""
import os
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn

BASE = os.path.dirname(os.path.abspath(__file__))
AI = os.path.dirname(BASE)
FIG = os.path.join(BASE, "figures")
DOCX = os.path.join(BASE, "张哲铭TASK8.docx")
FONT = "宋体"; SIZE = Pt(10.5)


def rfont(run, size=SIZE, bold=False, color=None, name=FONT, underline=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = size; run.font.bold = bold; run.font.underline = underline
    if color:
        run.font.color.rgb = color


def pfmt(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=0):
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(before); pf.space_after = Pt(after); pf.alignment = align


def body(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False, indent=True, size=SIZE, runs=None):
    p = doc.add_paragraph(); pfmt(p, align)
    if indent:
        p.paragraph_format.first_line_indent = Pt(21)
    if runs:
        for t, b, u in runs:
            rfont(p.add_run(t), size=size, bold=b, underline=u)
    else:
        rfont(p.add_run(text), size=size, bold=bold)
    return p


def h1(doc, text, size=Pt(14)):
    p = doc.add_paragraph(); pfmt(p, WD_ALIGN_PARAGRAPH.LEFT, before=6, after=2)
    rfont(p.add_run(text), size=size, bold=True); return p


def h2(doc, text, size=Pt(11.5)):
    p = doc.add_paragraph(); pfmt(p, WD_ALIGN_PARAGRAPH.LEFT, before=3)
    p.paragraph_format.first_line_indent = Pt(21)
    rfont(p.add_run(text), size=size, bold=True); return p


def caption(doc, text):
    p = doc.add_paragraph(); pfmt(p, WD_ALIGN_PARAGRAPH.CENTER)
    rfont(p.add_run(text), size=Pt(9), bold=True); return p


def pic(doc, name, width=6.1, base=FIG):
    path = os.path.join(base, name)
    if not os.path.exists(path):
        body(doc, f"[缺图 {name}]"); return
    p = doc.add_paragraph(); pfmt(p, WD_ALIGN_PARAGRAPH.CENTER)
    p.add_run().add_picture(path, width=Inches(width))


def make_table(doc, headers, rows, size=Pt(8.5)):
    table = doc.add_table(rows=1, cols=len(headers)); table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, hh in enumerate(headers):
        c = table.rows[0].cells[i]
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = c.paragraphs[0].paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE; pf.space_before = Pt(0); pf.space_after = Pt(0)
        rfont(c.paragraphs[0].add_run(hh), size=size, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = cells[i].paragraphs[0].paragraph_format
            pf.line_spacing_rule = WD_LINE_SPACING.SINGLE; pf.space_before = Pt(0); pf.space_after = Pt(0)
            rfont(cells[i].paragraphs[0].add_run(str(v)), size=size)
    return table


def P(x, d=1):
    try:
        return f"{float(x)*100:.{d}f}%"
    except Exception:
        return "-"


def F(x, d=2):
    try:
        return f"{float(x):.{d}f}"
    except Exception:
        return "-"


def toc_line(doc, text, page, level=0):
    p = doc.add_paragraph(); pfmt(p, WD_ALIGN_PARAGRAPH.LEFT)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    # 制表位右对齐 + 前导点
    from docx.shared import Cm
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(15.5), WD_TAB_ALIGNMENT.RIGHT, leader=2)  # leader=2 dotted
    indent = "　　" * level
    rfont(p.add_run(indent + text), size=SIZE, bold=(level == 0))
    rfont(p.add_run("\t" + str(page)), size=SIZE)
    return p


def build():
    pan = pd.read_csv(os.path.join(BASE, "panorama_summary.csv"), encoding="utf-8-sig")
    t7 = pd.read_csv(os.path.join(AI, "TASK7/data/compare_summary.csv"), encoding="utf-8-sig")
    def g(k, c):
        return t7[t7["策略"] == k][c].iloc[0]

    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    st.font.size = SIZE

    # ==================== 封面 ====================
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph(); pfmt(p, WD_ALIGN_PARAGRAPH.CENTER)
    rfont(p.add_run("量化交易策略开发与实践"), size=Pt(26), bold=True)
    p = doc.add_paragraph(); pfmt(p, WD_ALIGN_PARAGRAPH.CENTER)
    rfont(p.add_run("——北京大学 AI 量化工作坊学习成果综合报告"), size=Pt(15), bold=True)
    for _ in range(6):
        doc.add_paragraph()
    for label, val in [("作　　者：", "张哲铭"), ("所属项目：", "北京大学 AI 量化工作坊"),
                       ("报告主题：", "八项任务的策略开发、机器学习应用与实盘推演总结"),
                       ("完成日期：", "2026 年 7 月")]:
        p = doc.add_paragraph(); pfmt(p, WD_ALIGN_PARAGRAPH.CENTER)
        rfont(p.add_run(label + val), size=Pt(12), bold=False)
    doc.add_page_break()

    # ==================== 目录 ====================
    p = doc.add_paragraph(); pfmt(p, WD_ALIGN_PARAGRAPH.CENTER, after=6)
    rfont(p.add_run("目　录"), size=Pt(16), bold=True)
    toc = [
        ("摘要", "1", 0),
        ("一、量化交易核心概念", "2", 0),
        ("1.1 什么是量化交易", "2", 1),
        ("1.2 量化交易的核心价值", "2", 1),
        ("二、量化交易策略综合分析", "3", 0),
        ("2.1 八项任务与策略全景", "3", 1),
        ("2.2 各类策略的优缺点与适用场景", "4", 1),
        ("2.3 跨策略风险调整绩效对比", "5", 1),
        ("2.4 策略间的关联性与互补性", "6", 1),
        ("2.5 多策略量化交易系统的构建思路", "6", 1),
        ("三、机器学习在量化交易中的应用总结", "7", 0),
        ("3.1 数据预处理与特征工程", "7", 1),
        ("3.2 模型选择、训练与评估优化", "8", 1),
        ("3.3 机器学习的优势、局限与未来趋势", "8", 1),
        ("四、结论与展望", "9", 0),
        ("附录：改进建议", "10", 0),
    ]
    for t, pg, lv in toc:
        toc_line(doc, t, pg, lv)
    doc.add_page_break()

    # ==================== 摘要 ====================
    h1(doc, "摘要")
    body(doc,
        "本报告系统总结了本人在北京大学 AI 量化工作坊八项任务中的学习成果与实践经验。研究目的是打通"
        "“数据—指标—策略—机器学习—实盘—总结”的量化交易完整链条，并回答一个核心问题：在真实的 A 股市场"
        "（2018–2026 年、跨越多轮牛熊）中，哪一类策略能够以可控的风险获得稳健的超额收益。方法上，本人依次搭建了"
        "数据引擎、实现了经典技术指标，回测了双均线、海龟等趋势跟随策略，构建了基于机器学习的涨跌预测与择时策略，"
        "并最终在聚宽平台与本地环境中对小市值、银行股轮动、ETF 双重动量轮动三个策略完成了设计、优化与跨牛熊对比；"
        "全流程统一采用“信号次日成交、计单边万分之五成本、无未来函数”的严谨口径。主要成果是：以夏普比率与最大回撤"
        "为跨标的可比的评判标准，ETF 双重动量轮动策略以夏普 1.13、最大回撤 −21.3% 的表现居各类策略之首，且通过"
        "“调仓降频”这一反直觉改进实现了收益、夏普、回撤、成本的四重改善；银行股轮动与小市值策略在加入风控后分别"
        "成为“低回撤防御”与“稳健进攻”的代表。结论是：量化交易的核心竞争力不在于预测的“准”，而在于风险管理的“稳”；"
        "机器学习虽能提升信息利用效率，但在日频低信噪比场景下的绝对收益贡献有限，其价值更多体现在震荡下跌市的择时"
        "与控回撤上。报告最后就多策略组合、因子升级、执行优化等提出了改进建议（见附录）。")

    # ==================== 一、核心概念 ====================
    h1(doc, "一、量化交易核心概念")
    h2(doc, "1.1 什么是量化交易")
    body(doc,
        "量化交易是指借助数学模型、统计方法与计算机程序，将交易理念转化为明确、可执行、可回测的规则，"
        "并据此系统性地做出买卖决策的交易方式。它与传统手工交易的根本区别在于“纪律化”与“可验证”："
        "交易规则事先用代码固化，决策由数据与模型驱动，而非临场的情绪与主观判断。一套完整的量化交易系统通常包含"
        "数据引擎（获取并清洗行情、财务等数据）、因子与信号（从数据中提炼有预测力的特征）、策略逻辑（把信号转化为"
        "仓位）、回测框架（在历史数据上检验策略）与风控执行（止损、仓位管理、下单）五个环节——这恰好对应本工作坊"
        "八项任务的递进主线：任务一搭建数据引擎、任务二实现技术指标、任务三与任务四实现并回测经典策略、任务五与"
        "任务六引入机器学习、任务七完成平台实盘推演、任务八进行综合总结。")
    h2(doc, "1.2 量化交易的核心价值")
    body(doc,
        "综合八项任务的实践，本人将量化交易的核心价值归纳为四点。其一是纪律性——把规则写进程序，"
        "从根本上克服人性中的贪婪与恐惧，避免“追涨杀跌”。其二是可回测与可验证——任何想法都能在历史数据上量化检验，"
        "用夏普比率、最大回撤等客观指标评判优劣，而非凭感觉。其三是系统性与规模化——一套程序可同时监控成百上千个"
        "标的、执行复杂的多因子与轮动逻辑，这是人力无法企及的。其四是风险的可度量与可控制——量化框架天然支持止损、"
        "仓位管理与风险预算，把“亏多少”变成可以事先设定的参数。这四点中，本人体会最深的是最后一点：贯穿八项任务的"
        "最重要结论是，量化的核心竞争力不是把涨跌“预测得多准”，而是把风险“管理得多稳”。")

    # ==================== 二、策略综合分析 ====================
    h1(doc, "二、量化交易策略综合分析")
    h2(doc, "2.1 八项任务与策略全景")
    body(doc,
        "八项任务由浅入深地覆盖了量化交易的主要策略范式。为便于读者把握整体脉络，先以表 1 概览各任务的核心内容与"
        "策略类别，再逐类展开分析。")
    make_table(doc, ["任务", "主题", "策略/方法类别", "核心产出"],
        [["任务一", "数据引擎搭建", "基础设施", "行情数据获取、K线与收盘价曲线"],
         ["任务二", "技术指标实现", "技术分析", "MACD/RSI/KDJ/BOLL 等指标计算与解读"],
         ["任务三", "双均线策略", "趋势跟随", "金叉死叉信号，8 标的×3 周期回测"],
         ["任务四", "海龟交易法则", "通道突破趋势", "唐奇安通道+ATR 头寸管理，双系统回测"],
         ["任务五", "机器学习算法", "监督学习", "分类+回归多算法建模、评估与对比"],
         ["任务六", "机器学习定制策略", "ML 择时", "概率仓位+双阈值+风控的可回测策略"],
         ["任务七", "实盘推演", "动量/风格轮动", "小市值、银行股轮动、ETF 轮动三策略"],
         ["任务八", "成果总结", "综合分析", "本报告"]])
    caption(doc, "表 1　八项任务与策略全景一览")

    h2(doc, "2.2 各类策略的优缺点与适用场景")
    body(doc,
        "基于回测实证，本人将所涉策略归为四大范式，其优缺点与适用场景对比见表 2。")
    make_table(doc, ["策略范式", "代表任务", "优点", "缺点", "适用场景"],
        [["趋势跟随\n(双均线)", "任务三", "逻辑简单、大趋势中稳健", "震荡市反复被扫、参数敏感", "单边趋势明确的市场"],
         ["通道突破\n(海龟)", "任务四", "严格头寸/止损、纪律性强", "单边牛市易踏空、胜率偏低", "有大波段的品种"],
         ["机器学习\n择时", "任务五、六", "信息利用充分、可控风险", "日频信噪比低、易过拟合", "震荡下跌市控回撤"],
         ["动量/风格\n轮动", "任务七", "攻守兼备、分散风险", "拐点处动量崩溃、依赖标的池", "多资产、分化明显市场"]])
    caption(doc, "表 2　四大策略范式的优缺点与适用场景对比")
    body(doc,
        "从表 2 可见，趋势跟随与通道突破属于“单标的择时”，其共同短板是在单边上涨行情中因降低仓位而踏空，"
        "价值主要体现在控制回撤（任务三、四回测中，二者的绝对收益普遍跑输买入持有，但最大回撤显著更小）。"
        "机器学习择时（任务五、六）在日频方向预测上信噪比极低（AUC 仅约 0.5–0.6），因此其策略往往以极低仓位运行、"
        "绝对收益有限，但在下跌市中能有效规避损失。相比之下，动量与风格轮动（任务七）通过在多个标的间“选强汰弱”，"
        "在承担合理风险的前提下取得了最好的风险调整收益，是本工作坊中综合表现最优的一类。据此提出改进建议 1（见附录）。")

    h2(doc, "2.3 跨策略风险调整绩效对比")
    body(doc,
        "由于各任务的标的与回测区间不完全相同，直接比较绝对收益并不公平。本报告改用夏普比率（单位风险的超额收益）"
        "与最大回撤（最惨痛的资金损失）这两个跨标的可比的风险调整指标，对各类策略的代表性配置进行横向对比，"
        "结果如图 1 与表 3 所示。")
    pic(doc, "panorama.png", 6.4)
    caption(doc, "图 1　各类策略的风险调整绩效对比（A 夏普比率；B 最大回撤）")
    make_table(doc, ["策略类别", "代表配置", "夏普比率", "最大回撤"],
        [[r["策略类别"], r["代表配置"], F(r["夏普"]), P(r["最大回撤"])] for _, r in pan.iterrows()])
    caption(doc, "表 3　各类策略代表性配置的夏普比率与最大回撤")
    body(doc,
        "解读图 1 与表 3 需要一处关键说明：机器学习择时的夏普比率虽然名义上最高（约 1.70）、回撤极小（约 −2.3%），"
        "但这是因为该类策略平均仓位极低、大量时间处于空仓状态，波动与回撤自然被压得很小，其代价是绝对收益也非常"
        "微薄——这是“择时踏空”的另一种表现，属于“看起来很稳、实则赚得很少”。真正在充分投资状态下取得高夏普的是"
        "ETF 双重动量轮动策略（夏普 1.13、最大回撤 −21.3%），其含金量最高。银行股轮动（夏普 0.54、回撤仅 −20.5%）"
        "以三类策略中最小的回撤成为“防御担当”；小市值（夏普 0.40）经风控改造后回撤压至 −22.6%、各阶段均为正收益。"
        "作为对照，被动买入持有沪深300 夏普仅 0.27、最大回撤高达 −45.6%，全面弱于经过设计的主动策略。"
        "由此可见，优秀的策略设计确实能创造价值，据此提出改进建议 2（见附录）。")

    h2(doc, "2.4 策略间的关联性与互补性")
    body(doc,
        "各类策略并非彼此孤立，而是在“进攻—均衡—防御”的谱系上相互补位，这正是构建组合的基础。任务七的三个策略"
        "构成了一个天然的互补三角，其在不同牛熊阶段的表现对比如图 2 所示。")
    pic(doc, "task7_compare.png", 6.4)
    caption(doc, "图 2　三类主力策略的跨牛熊表现对比（净值、回撤、分阶段收益与风险-收益散点）")
    body(doc,
        "由图 2 可见：ETF 轮动攻守兼备、在牛市与反弹中弹性最强；小市值进攻性强、但需风控约束；"
        "银行股轮动依托高股息低波的防御属性、回撤最小，恰好在成长风格失效时提供保护。它们在不同市场阶段的表现"
        "此消彼长——例如在 2024 年四季度以来的成长股反弹中，ETF 轮动大幅领先而银行股轮动落后；但在 2022–2024 年的"
        "熊市里，银行股轮动逆势为正、成为组合的压舱石。趋势跟随与机器学习择时类策略则可作为“风险开关”，在系统性"
        "下跌中主动降低整体敞口。这种低相关、能互补的特性，是把它们组合成一个更稳健系统的前提，据此提出改进建议 3。")

    h2(doc, "2.5 多策略量化交易系统的构建思路")
    body(doc,
        "综合上述分析，本人提出一个分层的多策略量化交易系统构建思路。第一层是“核心配置层”，以攻守兼备的 ETF 双重"
        "动量轮动为主力，承担获取中期动量收益的主要职责。第二层是“卫星增强层”，配以进攻型的小市值策略与防御型的"
        "银行股轮动，通过风格分散在不同市场环境下平滑组合收益曲线。第三层是“风险控制层”，用趋势跟随/机器学习类的"
        "大盘择时信号作为“总闸”，在系统性风险来临时统一降低全组合仓位。三层之间按风险预算分配资金、定期再平衡，"
        "并对各子策略设置独立止损。这一“核心—卫星—风控”框架把单一策略的脆弱性分散到一个有机整体中，"
        "是本人对未来实盘的核心设计蓝图，具体落地要点见改进建议 3 与建议 4。")

    # ==================== 三、机器学习应用总结 ====================
    h1(doc, "三、机器学习在量化交易中的应用总结")
    body(doc,
        "任务五与任务六系统实践了机器学习在量化交易中的应用，本章按“数据预处理—特征工程—模型选择训练—评估优化”"
        "的流程总结关键要点与经验教训。")
    h2(doc, "3.1 数据预处理与特征工程")
    body(doc,
        "数据预处理的核心是保证“干净、对齐、无未来函数”。实践中最关键的三条经验是：其一，价格必须使用前复权数据，"
        "否则除权除息会造成价格跳空、污染收益计算；其二，财务等低频数据必须按“信息发布日”而非“报告期”对齐到日频"
        "特征上（任务五对个股财务因子即采用按发布日前向对齐的方法），否则会用到当时尚未公开的信息、造成回测虚高；"
        "其三，训练集与测试集必须严格按时间先后切分、绝不打乱，这是时间序列建模不可逾越的红线。特征工程方面，"
        "任务五共构造了约 38 个技术因子，涵盖动量/趋势、波动率、量价、经典技术指标（如相对强弱、平滑异同移动平均）"
        "与价格位置等大类，个股再叠加财务因子。这些自变量因子的设计遵循“有经济含义、无未来函数、稳定可计算”的原则。")
    h2(doc, "3.2 模型选择、训练与评估优化")
    body(doc,
        "在模型选择上，任务五对比了逻辑回归、决策树、随机森林、支持向量机、梯度提升树以及 XGBoost、LightGBM 等"
        "十余种分类与回归算法，用时间序列交叉验证做网格调参，以准确率、精确率、召回率、受试者工作特征曲线下面积"
        "（衡量分类排序能力的综合指标）等评估分类效果，以均方根误差、方向命中率等评估回归效果。核心发现有二："
        "一是树集成类模型（随机森林、梯度提升树系列）在多数标的上综合表现最优，与国际权威研究的结论一致；"
        "二是日频方向预测的信噪比极低，最优分类模型的曲线下面积也仅约 0.5–0.6、回归的拟合优度接近于零，"
        "这从根本上限制了策略的收益上限。任务六进一步把模型输出的“上涨概率”通过“双阈值降低换手、概率加权分配仓位、"
        "技术指标过滤、止损止盈”转化为可回测的交易策略，并为每个标的网格寻优。这一“预测—转化—风控”的完整流程，"
        "是机器学习真正落地为交易策略的关键，据此提出改进建议 5（见附录）。")
    h2(doc, "3.3 机器学习的优势、局限与未来趋势")
    body(doc,
        "机器学习的优势在于能处理高维、非线性的信息，自动挖掘传统线性规则难以刻画的特征交互，并以连续概率支持"
        "精细化的仓位管理。但其局限同样突出：一是过拟合风险高，样本内亮眼、样本外失效；二是高度依赖数据信噪比，"
        "在低信噪比的日频场景下上限很低；三是存在未来函数与信息泄漏的隐患，稍有不慎便前功尽弃；四是复杂模型可"
        "解释性差、且面临市场结构漂移导致的规律失效。因此，任务六中机器学习策略的绝对收益普遍不高、多在单边牛市中"
        "跑输买入持有，唯有在下跌市（如某互联网龙头个股的测试段）才显著跑赢——这生动印证了“机器学习择时的价值"
        "主要体现在震荡与下跌市的控回撤，而非单边牛市的收益增强”。展望未来，机器学习在量化中的发展趋势包括："
        "转向更长周期与截面排序（信噪比更高）、引入另类数据、采用深度学习做因子合成、以及“机器学习负责选股打分、"
        "规则策略负责风控执行”的人机结合范式。基于低信噪比的现实，本人提出改进建议 6（见附录）。")

    # ==================== 四、结论与展望 ====================
    h1(doc, "四、结论与展望")
    body(doc,
        "回顾八项任务的学习历程，本人在三个层面收获颇丰。在认知层面，建立了对量化交易完整链条的系统理解，"
        "深刻体会到“风险管理的稳比预测的准更重要”这一贯穿始终的核心理念。在技术层面，掌握了从数据引擎搭建、"
        "技术指标实现、经典策略回测，到机器学习建模、策略优化与平台实盘推演的全套技能，并养成了“次日成交、计成本、"
        "无未来函数”的严谨习惯。在实践层面，通过对多类策略在跨牛熊数据上的亲手回测与对比，积累了大量真实、可复用的"
        "经验——尤其是“适度降低调仓频率可同时改善收益与风险”“冗余的择时风控反而会侵蚀超额收益”“高相关板块控制"
        "系统性风险比板块内选股更重要”等反直觉但极具价值的结论。")
    body(doc,
        "展望未来，本人的进一步探索方向有三：一是落地附录所列的多策略组合系统，用真实资金做小规模实盘验证；"
        "二是升级因子体系，引入基本面景气度、资金流与另类数据，并尝试用深度学习做因子合成；三是完善执行与风控细节，"
        "包括交易成本与冲击成本的精细建模、动态风险预算与组合再平衡机制。量化交易是一条“持续迭代、敬畏市场”的长路，"
        "本次工作坊为本人打下了坚实的方法论与实践基础，未来将继续在真实市场的检验中不断优化与成长。")

    # ==================== 附录：改进建议 ====================
    doc.add_page_break()
    h1(doc, "附录：改进建议")
    body(doc, "以下改进建议根据正文的分析与推断得出，逐项编号，正文相应位置已按编号引用。", indent=True)
    adv = [
        ("建议 1", "以动量/风格轮动类策略为核心。 回测证明其风险调整收益最优，应作为未来策略研发与资金配置的重点，"
                   "而非把主要精力放在收益上限受限的单标的日频择时上。"),
        ("建议 2", "始终以风险调整指标评判策略、警惕“低仓位高夏普”的假象。 评估时应同时考察夏普比率、最大回撤与"
                   "平均仓位/资金利用率，避免被“空仓换来的高夏普低回撤”误导，确保收益是在充分投资下取得的。"),
        ("建议 3", "构建低相关的多策略组合。 将攻守兼备的 ETF 轮动、进攻型小市值、防御型银行股轮动按风险预算组合，"
                   "利用其在不同市场阶段的互补性平滑收益曲线，显著优于单押任一策略。"),
        ("建议 4", "搭建“核心—卫星—风控”三层系统并引入总仓位择时。 以趋势/机器学习类大盘信号作为组合级“风险总闸”，"
                   "在系统性下跌时统一降低敞口；各子策略保留独立止损，定期再平衡。"),
        ("建议 5", "坚持“预测—转化—风控”的机器学习落地范式并严防未来函数。 特征按发布日对齐、训练测试按时间切分、"
                   "结果次日成交、计入成本，是一切结论可信的前提；宁可结论“不好看”，也要真实。"),
        ("建议 6", "把机器学习的应用从日频方向预测转向更高信噪比场景。 优先用于中长周期、横截面排序选股与因子合成，"
                   "并采用“机器学习选股打分 + 规则策略风控执行”的人机结合模式，扬长避短。"),
    ]
    for tag, text in adv:
        p = doc.add_paragraph(); pfmt(p, WD_ALIGN_PARAGRAPH.JUSTIFY)
        p.paragraph_format.first_line_indent = Pt(21)
        rfont(p.add_run(tag + "："), bold=True)
        rfont(p.add_run(text))

    doc.add_paragraph()
    body(doc, "（注：本报告为量化学习实践总结，所用为历史数据回测，不构成任何投资建议；市场有风险，决策需谨慎。"
              "　作者：张哲铭）", indent=True, size=Pt(9))

    doc.save(DOCX)
    print(f"已生成 {DOCX}")
    return DOCX


if __name__ == "__main__":
    build()
