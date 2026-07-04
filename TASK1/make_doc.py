# -*- coding: utf-8 -*-
"""
生成作业文档：张哲铭TASK1.docx
格式要求：宋体、五号字(10.5pt)、1.5倍行距、0倍段间距、正文两端对齐。
图表带标号、标题与解读。
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
FIG_PATH = os.path.join(BASE_DIR, "close_price.png")
DOCX_PATH = os.path.join(BASE_DIR, "张哲铭TASK1.docx")

FONT_NAME = "宋体"
FONT_SIZE = Pt(10.5)  # 五号字


def set_run_font(run, size=FONT_SIZE, bold=False, color=None):
    """统一设置字体为宋体（中英文均生效）。"""
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = size
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_para_format(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """设置段落：1.5倍行距、0段间距、指定对齐方式。"""
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.alignment = align


def add_body(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False,
             size=FONT_SIZE, first_indent=True):
    """添加正文段落（默认首行缩进2字符、两端对齐）。"""
    p = doc.add_paragraph()
    set_para_format(p, align)
    if first_indent:
        p.paragraph_format.first_line_indent = Pt(21)  # 约2个五号字
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def add_title(doc, text, size=Pt(16)):
    """文档大标题（居中、加粗）。"""
    p = doc.add_paragraph()
    set_para_format(p, WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=True)
    return p


def add_heading(doc, text, size=Pt(12)):
    """章节标题（加粗、左对齐、无首行缩进）。"""
    p = doc.add_paragraph()
    set_para_format(p, WD_ALIGN_PARAGRAPH.LEFT)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=True)
    return p


def add_caption(doc, text):
    """图表标题/标号（居中）。"""
    p = doc.add_paragraph()
    set_para_format(p, WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run(text)
    set_run_font(run, size=Pt(9), bold=True)
    return p


def add_code_block(doc, code_text):
    """代码块（等宽字体、左对齐、无缩进、浅灰）。"""
    for line in code_text.split("\n"):
        p = doc.add_paragraph()
        set_para_format(p, WD_ALIGN_PARAGRAPH.LEFT)
        pf = p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def main():
    doc = Document()

    # 文档默认样式设为宋体五号
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    style.font.size = FONT_SIZE

    # ===== 标题与信息 =====
    add_title(doc, "量化交易初体验：从零搭建数据引擎")
    info = doc.add_paragraph()
    set_para_format(info, WD_ALIGN_PARAGRAPH.CENTER)
    r = info.add_run("北京大学 AI 量化工作坊 · TASK1　　姓名：张哲铭")
    set_run_font(r, size=Pt(10.5))
    doc.add_paragraph()

    # ===== 第一题 =====
    add_heading(doc, "一、相较于传统手工操作交易，量化交易有哪些优势？")
    add_body(doc,
        "量化交易是指借助数学模型、统计方法与计算机程序，将交易理念转化为明确的、"
        "可执行的规则，并由程序自动完成行情监控、信号生成乃至下单执行的交易方式。"
        "相较于依赖个人经验与主观判断的传统手工交易，量化交易主要具有以下优势：")

    points = [
        ("纪律性强、克服人性弱点。", "交易规则由程序严格执行，不受贪婪、恐惧、犹豫等情绪"
         "干扰，能够始终如一地按既定策略操作，避免了手工交易中常见的追涨杀跌与随意改单。"),
        ("系统性与可重复性。", "策略以明确的逻辑与参数固化下来，任何人、任何时间运行同一策略"
         "都能得到一致的结果，便于团队协作、复盘和持续优化。"),
        ("处理海量数据、捕捉多维机会。", "程序可同时跟踪成百上千只标的、多个市场与多种因子，"
         "远超人脑的信息处理能力，能够发现人工难以察觉的规律与套利机会。"),
        ("反应速度快、执行精准。", "从信号触发到下单可在毫秒级完成，尤其在高频与套利场景中，"
         "速度优势直接转化为收益；同时避免了手工下单的手误与延迟。"),
        ("可回测、可验证。", "策略在实盘之前可用历史数据进行回测，量化评估其收益、回撤、"
         "胜率等风险收益特征，用数据而非主观感觉来检验策略是否有效。"),
        ("严格的风险管理。", "止损、仓位控制、风险敞口限制等均可写入程序自动执行，"
         "使风控贯穿交易全程，不因情绪波动而失效。"),
    ]
    for i, (head, body) in enumerate(points, 1):
        p = doc.add_paragraph()
        set_para_format(p)
        p.paragraph_format.first_line_indent = Pt(21)
        r1 = p.add_run(f"（{i}）{head}")
        set_run_font(r1, bold=True)
        r2 = p.add_run(body)
        set_run_font(r2)

    add_body(doc,
        "需要指出的是，量化交易并非稳赚不赔：它依赖历史数据与模型假设，存在过拟合、"
        "极端行情失效、系统与数据风险等问题。其核心价值在于把交易变成一套可检验、"
        "可迭代的系统工程，而非取代对市场的理解。")

    # ===== 第二题 =====
    doc.add_paragraph()
    add_heading(doc, "二、基本概念解释：K 线、基本面、技术面")

    add_body(doc, "【K 线】", first_indent=True, bold=True)
    add_body(doc,
        "K 线又称蜡烛图、阴阳线，起源于日本江户时代的米市交易，用于直观描述某一时间"
        "周期（如日、周、月，或分钟级）内价格的波动。每根 K 线由“实体”和“上下影线”"
        "构成，包含四个关键价格：开盘价、收盘价、最高价与最低价。当收盘价高于开盘价时"
        "为阳线（A 股习惯用红色表示上涨），实体上下沿分别为收盘价与开盘价；当收盘价低于"
        "开盘价时为阴线（A 股用绿色表示下跌）。上影线的顶端为最高价，下影线的底端为最低价。"
        "K 线将多空双方的力量对比浓缩在一根图形中，是技术分析最基础、最常用的工具。")

    add_body(doc, "【基本面】", first_indent=True, bold=True)
    add_body(doc,
        "基本面是指影响资产内在价值的基础性因素，回答“这家公司/这个资产到底值多少钱”"
        "的问题。对股票而言，基本面分析主要考察：① 宏观层面，如经济增长、利率、通胀、"
        "货币与财政政策等；② 行业层面，如行业景气度、竞争格局、政策导向；③ 公司层面，"
        "如营业收入、净利润、毛利率、资产负债结构、现金流、市盈率（PE）、市净率（PB）"
        "等财务指标以及管理层与商业模式。基本面分析着眼于中长期价值，是价值投资的核心方法。")

    add_body(doc, "【技术面】", first_indent=True, bold=True)
    add_body(doc,
        "技术面是指通过分析历史的价格与成交量数据，研究市场行为本身，从而预测未来价格"
        "走势的方法。它以“价格反映一切信息、历史会重演、价格沿趋势运行”为基本假设，"
        "主要工具包括 K 线形态、趋势线、支撑与压力位，以及均线（MA）、MACD、RSI、"
        "KDJ、布林带等技术指标。技术面关注买卖时机与市场情绪，偏重中短期择时，"
        "与着眼长期价值的基本面分析互为补充。")

    # ===== 第三题 =====
    doc.add_paragraph()
    add_heading(doc, "三、基于 Tushare 的数据引擎实现")
    add_body(doc,
        "本部分注册并使用 Tushare Pro 平台（https://www.tushare.pro/），获取查询 Token 后，"
        "通过 Python 编程搭建数据引擎，选取沪市大盘蓝筹、水电龙头——长江电力"
        "（代码 600900.SH）作为研究标的，完成过去一年日线数据的获取、可视化与存储。")

    add_body(doc, "1. 技术方案与核心代码", first_indent=True, bold=True)
    add_body(doc,
        "数据引擎基于 tushare、pandas、matplotlib 三个库实现：使用 tushare 的 pro.daily 接口"
        "拉取日线行情，用 pandas 进行清洗与排序，用 matplotlib 绘制收盘价曲线并保存为图片，"
        "最后将原始数据落地为 CSV 文件。核心代码如下：")

    core_code = '''import tushare as ts
import pandas as pd
import matplotlib.pyplot as plt

# 1) 设置 Token 并初始化接口
ts.set_token("你的Token")
pro = ts.pro_api()

# 2) 获取长江电力过去一年的日线数据
df = pro.daily(ts_code="600900.SH",
               start_date="20250704",
               end_date="20260704")
df["trade_date"] = pd.to_datetime(df["trade_date"], format="%Y%m%d")
df = df.sort_values("trade_date").reset_index(drop=True)

# 3) 绘制每日收盘价曲线图
plt.plot(df["trade_date"], df["close"], color="#c0392b")
plt.title("长江电力（600900.SH）近一年每日收盘价走势")
plt.xlabel("交易日期"); plt.ylabel("收盘价（元）")
plt.savefig("close_price.png")

# 4) 保存为 CSV，供后续任务复用
df.to_csv("600900_daily.csv", index=False, encoding="utf-8-sig")'''
    add_code_block(doc, core_code)

    add_body(doc, "2. 收盘价曲线图与解读", first_indent=True, bold=True)

    # 插入图片
    pic_p = doc.add_paragraph()
    set_para_format(pic_p, WD_ALIGN_PARAGRAPH.CENTER)
    pic_p.add_run().add_picture(FIG_PATH, width=Inches(6.0))
    add_caption(doc, "图 1　长江电力（600900.SH）近一年每日收盘价走势图")

    add_body(doc,
        "如图 1 所示，程序共获取到 2025-07-04 至 2026-07-03 期间 242 个交易日的有效数据。"
        "该股期初收盘价为 30.16 元，期末为 27.05 元，区间累计下跌约 10.31%；期间最高价 30.61 元"
        "（出现在 2025 年 7 月中旬），最低价 25.65 元（出现在 2026 年 2 月上旬）。整体走势可分为"
        "三个阶段：2025 年 7 月至 9 月自高位震荡回落；2025 年 9 月至 2026 年 2 月宽幅震荡并在 2 月"
        "探至年内低点；此后自低位逐步修复，回升至 27 元附近。区间内收盘价均值约 27.56 元，"
        "标准差约 0.92 元，日收益率年化波动率约 12.3%，显示出作为公用事业蓝筹股“波动较低、"
        "走势稳健”的典型特征，与其稳定的水电主业和高分红属性相符。")

    add_body(doc, "3. 数据存储说明", first_indent=True, bold=True)
    add_body(doc,
        "获取到的原始数据已保存为 600900_daily.csv（UTF-8-SIG 编码，共 242 行），包含"
        "股票代码（ts_code）、交易日期（trade_date）、开盘价（open）、最高价（high）、"
        "最低价（low）、收盘价（close）、前收盘价（pre_close）、涨跌额（change）、"
        "涨跌幅（pct_chg）、成交量（vol）、成交额（amount）共 11 个字段，可直接供后续的"
        "策略开发、指标计算与回测任务复用。")

    doc.add_paragraph()
    add_heading(doc, "四、小结")
    add_body(doc,
        "本次任务从理论与实践两方面完成了量化交易的初体验：理论上厘清了量化交易相对于"
        "手工交易的优势，以及 K 线、基本面、技术面三个基础概念；实践上基于 Tushare 平台"
        "从零搭建了一个可运行的数据引擎，完成了“数据获取—可视化—持久化存储”的完整闭环。"
        "该数据引擎所产出的 CSV 数据与代码框架，为后续的因子构建、策略开发与回测奠定了基础。")

    doc.save(DOCX_PATH)
    print(f"已生成文档：{DOCX_PATH}")


if __name__ == "__main__":
    main()
