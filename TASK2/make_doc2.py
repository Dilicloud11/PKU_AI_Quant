# -*- coding: utf-8 -*-
"""
生成 TASK2 作业文档：张哲铭TASK2.docx
格式：宋体、五号(10.5pt)、1.5倍行距、0段间距、正文两端对齐。
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

BASE = os.path.dirname(os.path.abspath(__file__))
DOCX = os.path.join(BASE, "张哲铭TASK2.docx")
FONT = "宋体"
SIZE = Pt(10.5)


def rfont(run, size=SIZE, bold=False, color=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = size
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def pfmt(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.alignment = align


def body(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False, indent=True, size=SIZE):
    p = doc.add_paragraph()
    pfmt(p, align)
    if indent:
        p.paragraph_format.first_line_indent = Pt(21)
    r = p.add_run(text)
    rfont(r, size=size, bold=bold)
    return p


def title(doc, text, size=Pt(16)):
    p = doc.add_paragraph(); pfmt(p, WD_ALIGN_PARAGRAPH.CENTER)
    rfont(p.add_run(text), size=size, bold=True); return p


def heading(doc, text, size=Pt(13)):
    p = doc.add_paragraph(); pfmt(p, WD_ALIGN_PARAGRAPH.LEFT)
    rfont(p.add_run(text), size=size, bold=True); return p


def subheading(doc, text, size=Pt(11)):
    p = doc.add_paragraph(); pfmt(p, WD_ALIGN_PARAGRAPH.LEFT)
    p.paragraph_format.first_line_indent = Pt(21)
    rfont(p.add_run(text), size=size, bold=True); return p


def caption(doc, text):
    p = doc.add_paragraph(); pfmt(p, WD_ALIGN_PARAGRAPH.CENTER)
    rfont(p.add_run(text), size=Pt(9), bold=True); return p


def pic(doc, name, width=6.1):
    p = doc.add_paragraph(); pfmt(p, WD_ALIGN_PARAGRAPH.CENTER)
    p.add_run().add_picture(os.path.join(BASE, name), width=Inches(width))


def formula(doc, text):
    """公式/计算式：居中、楷体感，用较小字号突出。"""
    p = doc.add_paragraph(); pfmt(p, WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run(text)
    r.font.name = "Cambria Math"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(10.5)
    r.font.italic = True


def code_block(doc, code):
    for line in code.split("\n"):
        p = doc.add_paragraph(); pfmt(p, WD_ALIGN_PARAGRAPH.LEFT)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        r = p.add_run(line if line else " ")
        r.font.name = "Consolas"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def main():
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    st.font.size = SIZE

    title(doc, "数据炼金术：数据诊断与构造交易指标")
    info = doc.add_paragraph(); pfmt(info, WD_ALIGN_PARAGRAPH.CENTER)
    rfont(info.add_run("北京大学 AI 量化工作坊 · TASK2　　姓名：张哲铭"), size=Pt(10.5))
    doc.add_paragraph()

    body(doc,
        "本任务承接 TASK1 搭建的数据引擎，以长江电力（600900.SH）近一年（2025-07-04 至 "
        "2026-07-03，共 242 个交易日）的日线数据为对象，先对数据进行基础诊断，再构造 RSI、"
        "MACD、布林带三个主流技术指标并可视化，最后扩展介绍并实现 KDJ 指标。")

    # ===== 一、数据诊断 =====
    heading(doc, "一、数据基础诊断分析")
    body(doc,
        "数据诊断是量化分析的第一道关口，目的是在计算指标、开发策略之前确认数据的完整性与"
        "合理性，避免“垃圾进、垃圾出”。本部分从缺失值检查、重复值检查与描述性统计三方面展开。")

    subheading(doc, "1. 缺失值与重复值检查")
    body(doc,
        "使用 pandas 的 isnull().sum() 逐字段统计缺失数量，并用 duplicated() 检查重复交易日。"
        "检查结果表明：全部 11 个字段、242 行数据均无缺失值，且无重复交易日，数据完整、"
        "可直接用于后续指标计算。核心代码如下：")
    code_block(doc,
        'df.isnull().sum()                 # 各字段缺失值数量\n'
        'df.duplicated(subset=["trade_date"]).sum()   # 重复交易日\n'
        'df.describe()                     # 描述性统计量')

    subheading(doc, "2. 描述性统计量")
    body(doc,
        "对开盘价、最高价、最低价、收盘价、成交量、成交额与涨跌幅计算描述性统计量，"
        "结果见表 1。可以看出：区间内收盘价均值约 27.56 元、标准差仅 0.92 元，最低 25.65 元、"
        "最高 30.61 元，价格波动区间较窄；日涨跌幅均值约 -0.03%、标准差 0.76%，且最大单日涨幅"
        "（+2.05%）与最大单日跌幅（-3.00%）都不大，整体呈现大盘蓝筹“低波动、走势稳健”的特征。"
        "成交量与成交额的标准差相对均值较大，说明市场活跃度在不同交易日之间存在明显差异。")

    # 表1：描述性统计
    import pandas as pd
    desc = pd.read_csv(os.path.join(BASE, "describe_stats.csv"), index_col=0)
    cap_map = {"open":"开盘价","high":"最高价","low":"最低价","close":"收盘价",
               "vol":"成交量","amount":"成交额","pct_chg":"涨跌幅(%)"}
    caption(doc, "表 1　核心字段描述性统计量")
    cols = ["mean","std","min","25%","50%","75%","max"]
    head = ["字段","均值","标准差","最小值","25%分位","中位数","75%分位","最大值"]
    table = doc.add_table(rows=1, cols=len(head)); table.style = "Table Grid"
    for i, h in enumerate(head):
        c = table.rows[0].cells[i]; c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        rfont(c.paragraphs[0].add_run(h), size=Pt(9), bold=True)
    for idx, row in desc.iterrows():
        cells = table.add_row().cells
        vals = [cap_map.get(idx, idx)]
        for cc in cols:
            v = row[cc]
            vals.append(f"{v:,.2f}" if abs(v) >= 100 else f"{v:.3f}")
        for i, v in enumerate(vals):
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            rfont(cells[i].paragraphs[0].add_run(str(v)), size=Pt(8.5))

    # ===== 二、三大指标原理 =====
    doc.add_paragraph()
    heading(doc, "二、RSI、MACD、布林带的计算方法与作用")

    subheading(doc, "1. RSI（相对强弱指标，Relative Strength Index）")
    body(doc,
        "RSI 由 Welles Wilder 提出，通过比较一段时期内价格上涨与下跌的力量强弱，衡量市场"
        "多空双方的相对强度，取值范围 0~100。计算方法（以 N=14 日为例）：首先计算每日价格变动，"
        "分别求出 N 日内的平均涨幅与平均跌幅，二者之比记为 RS，再代入下式：")
    formula(doc, "RS = 平均涨幅 / 平均跌幅 ，   RSI = 100 − 100 / (1 + RS)")
    body(doc,
        "作用：RSI 主要用于判断超买超卖。通常 RSI > 70 视为超买（价格可能回落），RSI < 30 "
        "视为超卖（价格可能反弹）；50 为多空分界。此外，RSI 与价格的背离（价格创新高而 RSI 未"
        "创新高）常被视为趋势反转的预警信号。")

    subheading(doc, "2. MACD（指数平滑异同移动平均线，Moving Average Convergence Divergence）")
    body(doc,
        "MACD 由 Gerald Appel 提出，基于快、慢两条指数移动平均线（EMA）的差离来刻画趋势"
        "的方向与动能。以常用参数 (12, 26, 9) 为例，计算步骤为：")
    formula(doc, "DIF = EMA(12) − EMA(26)")
    formula(doc, "DEA = EMA(DIF, 9)")
    formula(doc, "MACD 柱 = 2 × (DIF − DEA)")
    body(doc,
        "作用：DIF 上穿 DEA 形成“金叉”，为买入信号；DIF 下穿 DEA 形成“死叉”，为卖出信号。"
        "MACD 柱（红柱为正、绿柱为负）反映两线差距的变化，柱体由绿翻红、由红翻绿分别预示动能"
        "转强或转弱。MACD 兼具趋势跟踪与动能判断能力，是应用最广泛的中线指标之一。")

    subheading(doc, "3. 布林带（Bollinger Bands）")
    body(doc,
        "布林带由 John Bollinger 提出，以移动平均线为中枢、以价格标准差刻画波动区间，由中轨、"
        "上轨、下轨三条线组成。以常用参数 (20, 2) 为例：")
    formula(doc, "中轨 = MA(20)")
    formula(doc, "上轨 = 中轨 + 2 × σ(20) ，  下轨 = 中轨 − 2 × σ(20)")
    body(doc,
        "其中 σ(20) 为收盘价 20 日标准差。作用：在统计意义上，价格约 95% 的时间运行于上下轨"
        "之间。价格触及上轨往往意味着短期偏强或超买，触及下轨意味着偏弱或超卖；轨道“收窄”"
        "预示波动率降低、变盘临近，轨道“张口”则表示趋势与波动放大。布林带把趋势与波动率"
        "融为一体，常用于判断相对高低位与突破。")

    # ===== 三、编程实现与可视化 =====
    doc.add_paragraph()
    heading(doc, "三、Python 编程实现与指标可视化")
    body(doc,
        "本部分加载 TASK1 存储的 600900_daily.csv，手工实现上述三个指标（不依赖第三方指标库，"
        "以便清晰呈现计算逻辑），并绘制可视化图形。三个指标的核心实现代码如下：")
    code_block(doc,
        '# RSI(14)：用 Wilder 平滑\n'
        'delta = close.diff()\n'
        'gain = delta.clip(lower=0); loss = -delta.clip(upper=0)\n'
        'avg_gain = gain.ewm(alpha=1/14, adjust=False).mean()\n'
        'avg_loss = loss.ewm(alpha=1/14, adjust=False).mean()\n'
        'rsi = 100 - 100 / (1 + avg_gain/avg_loss)\n\n'
        '# MACD(12,26,9)\n'
        'dif = close.ewm(span=12).mean() - close.ewm(span=26).mean()\n'
        'dea = dif.ewm(span=9).mean()\n'
        'hist = (dif - dea) * 2\n\n'
        '# 布林带(20,2)\n'
        'mid = close.rolling(20).mean(); std = close.rolling(20).std()\n'
        'upper = mid + 2*std; lower = mid - 2*std')

    subheading(doc, "1. RSI 指标可视化")
    pic(doc, "rsi.png")
    caption(doc, "图 1　长江电力收盘价与 RSI(14) 指标")
    body(doc,
        "如图 1，样本期内 RSI 全程未突破 70 超买线（最高约 69.3），却有 17 个交易日跌破 30 "
        "超卖线（最低约 21.3），主要集中在 2026 年 1—2 月股价探底阶段——这与该股当时创出年内"
        "低点 25.65 元相互印证，超卖信号出现后价格随即企稳回升，显示 RSI 对短期底部具有一定"
        "提示作用。整体 RSI 多在 30~65 区间运行，符合震荡偏弱行情的特征。")

    subheading(doc, "2. MACD 指标可视化")
    pic(doc, "macd.png")
    caption(doc, "图 2　长江电力收盘价与 MACD(12,26,9) 指标")
    body(doc,
        "如图 2，样本期内 DIF 与 DEA 共发生约 7 次金叉与 7 次死叉，金叉多出现在阶段性底部"
        "（如 2025 年 9 月、2026 年 3 月），死叉多出现在阶段性顶部，与股价的波段起落基本吻合。"
        "MACD 柱由绿翻红对应动能转强、由红翻绿对应动能转弱。需要注意的是，在 2025 年 10—12 月"
        "的窄幅震荡中，MACD 出现较多小幅金叉死叉，属于震荡行情中的“假信号”，提示该指标在"
        "无趋势市中需结合其他工具过滤噪音。")

    subheading(doc, "3. 布林带可视化")
    pic(doc, "boll.png")
    caption(doc, "图 3　长江电力布林带 Bollinger Bands(20,2)")
    body(doc,
        "如图 3，收盘价绝大多数时间运行于上下轨构成的通道内，与“约 95% 时间处于带内”的统计"
        "规律一致。样本期内价格触及或突破上轨约 6 次（多为阶段高点），触及或跌破下轨约 17 次"
        "（多为阶段低点，集中于 2026 年初的下跌段）。2026 年 2 月股价跌破下轨后快速收回，是典型的"
        "超跌反弹；而在 2025 年 12 月至 2026 年 1 月，可见带宽明显收窄后随即张口下行，直观展示了"
        "“先收口、后变盘”的波动率变化过程。")

    # ===== 四、扩展指标 KDJ =====
    doc.add_paragraph()
    heading(doc, "四、扩展指标：KDJ 随机指标")
    body(doc,
        "除上述三个指标外，常见的技术指标还包括 KDJ（随机指标）、DMA、OBV（能量潮）、"
        "威廉指标 %R、乖离率 BIAS、ATR（真实波动幅度）、CCI（顺势指标）等。本部分选取应用"
        "十分广泛的 KDJ 指标进行介绍与实现。")
    body(doc,
        "KDJ 由 George Lane 提出，综合了动量、强弱与移动平均的思想，对价格在近期高低区间中"
        "所处的位置更为敏感。以参数 (9, 3, 3) 为例，计算步骤为：先求 N=9 日内的未成熟随机值 RSV，"
        "再对其平滑得到 K、D，并由 K、D 线性组合得到 J：")
    formula(doc, "RSV = (收盘价 − 9日最低价) / (9日最高价 − 9日最低价) × 100")
    formula(doc, "K = SMA(RSV, 3) ， D = SMA(K, 3) ， J = 3K − 2D")
    body(doc,
        "作用：KDJ 主要用于判断超买超卖与捕捉转折。通常 K、D 在 80 以上为超买区、20 以下为"
        "超卖区；K 上穿 D 为金叉（买入），K 下穿 D 为死叉（卖出）。J 值波动最快、最灵敏，常作为"
        "领先信号。KDJ 在震荡行情中表现尤佳，但在单边强趋势中易出现“钝化”。实现代码如下：")
    code_block(doc,
        'low_n = df["low"].rolling(9).min()\n'
        'high_n = df["high"].rolling(9).max()\n'
        'rsv = (df["close"] - low_n) / (high_n - low_n) * 100\n'
        'k = rsv.ewm(alpha=1/3, adjust=False).mean()\n'
        'd = k.ewm(alpha=1/3, adjust=False).mean()\n'
        'j = 3*k - 2*d')
    pic(doc, "kdj.png")
    caption(doc, "图 4　长江电力收盘价与 KDJ(9,3,3) 指标")
    body(doc,
        "如图 4，样本期内 K 值有 13 个交易日进入 80 以上超买区、40 个交易日跌入 20 以下超卖区，"
        "与该股震荡偏弱、底部区域停留时间较长的走势相符。J 线波动幅度最大，频繁在 0~100 之外"
        "穿刺，灵敏地领先于 K、D 反映短期拐点。K 与 D 的金叉、死叉与股价的短线波段有较好的"
        "对应关系，验证了 KDJ 在震荡市中较强的择时能力。")

    # ===== 小结 =====
    doc.add_paragraph()
    heading(doc, "五、小结")
    body(doc,
        "本次任务完成了从数据诊断到指标构造的完整“数据炼金”流程：先确认了数据的完整性与"
        "统计特征，再从原理、计算方法到编程实现，系统地构建并可视化了 RSI、MACD、布林带三大"
        "主流指标，并扩展实现了 KDJ 指标。四个指标从不同角度（超买超卖、趋势动能、波动区间、"
        "随机强弱）刻画了长江电力近一年的市场状态，彼此印证又各有侧重。这些指标为后续的交易"
        "信号生成与策略回测提供了直接可用的特征基础。")

    doc.save(DOCX)
    print("已生成：", DOCX)


if __name__ == "__main__":
    main()
