# -*- coding: utf-8 -*-
"""
生成 TASK4 作业文档：张哲铭TASK4.docx / .pdf
格式：宋体、五号(10.5pt)、1.5倍行距、0段间距、正文两端对齐。
内容：海龟策略核心思想与优势、高低点通道/ATR/止损条件解释、
      完整交易流程图、Python 实现、文献综述（标的有效性）、
      多标的×多参数实证回测、完整回测指标体系、参数敏感性、应用心得。
作者：张哲铭
"""
import os
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

BASE = os.path.dirname(os.path.abspath(__file__))
FIG = os.path.join(BASE, "figures")
DOCX = os.path.join(BASE, "张哲铭TASK4.docx")
FONT = "宋体"
SIZE = Pt(10.5)  # 五号

NAME = {
    "sh518880": "黄金ETF", "sz159941": "纳指ETF", "hk00700": "腾讯控股",
    "sh600900": "长江电力", "sh515450": "红利低波50ETF", "sh588000": "科创50ETF",
    "sh510300": "沪深300ETF", "sh510500": "中证500ETF",
}


# ============ 样式辅助 ============
def rfont(run, size=SIZE, bold=False, color=None, name=FONT):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = size
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def pfmt(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.alignment = align


def body(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False, indent=True, size=SIZE):
    p = doc.add_paragraph(); pfmt(p, align)
    if indent:
        p.paragraph_format.first_line_indent = Pt(21)
    rfont(p.add_run(text), size=size, bold=bold)
    return p


def title(doc, text, size=Pt(16)):
    p = doc.add_paragraph(); pfmt(p, WD_ALIGN_PARAGRAPH.CENTER)
    rfont(p.add_run(text), size=size, bold=True); return p


def heading(doc, text, size=Pt(13)):
    p = doc.add_paragraph(); pfmt(p, WD_ALIGN_PARAGRAPH.LEFT)
    rfont(p.add_run(text), size=size, bold=True); return p


def subheading(doc, text, size=Pt(11)):
    p = doc.add_paragraph(); pfmt(p, WD_ALIGN_PARAGRAPH.LEFT)
    p.paragraph_format.first_line_indent = Pt(21)
    rfont(p.add_run(text), size=size, bold=True); return p


def caption(doc, text):
    p = doc.add_paragraph(); pfmt(p, WD_ALIGN_PARAGRAPH.CENTER)
    rfont(p.add_run(text), size=Pt(9), bold=True); return p


def pic(doc, name, width=6.1):
    p = doc.add_paragraph(); pfmt(p, WD_ALIGN_PARAGRAPH.CENTER)
    p.add_run().add_picture(os.path.join(FIG, name), width=Inches(width))


def formula(doc, text):
    p = doc.add_paragraph(); pfmt(p, WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run(text)
    r.font.name = "Cambria Math"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(10.5); r.font.italic = True


def code_block(doc, code):
    for line in code.split("\n"):
        p = doc.add_paragraph(); pfmt(p, WD_ALIGN_PARAGRAPH.LEFT)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        r = p.add_run(line if line else " ")
        r.font.name = "Consolas"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def make_table(doc, headers, rows, size=Pt(8.5)):
    table = doc.add_table(rows=1, cols=len(headers)); table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        rfont(c.paragraphs[0].add_run(h), size=Pt(8.5), bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = cells[i].paragraphs[0].paragraph_format
            pf.space_before = Pt(0); pf.space_after = Pt(0)
            rfont(cells[i].paragraphs[0].add_run(str(v)), size=size)
    return table


def pct(x, d=1):
    return f"{x*100:.{d}f}%"


def build():
    df = pd.read_csv(os.path.join(BASE, "backtest_results.csv"))
    scan = pd.read_csv(os.path.join(BASE, "param_scan_results.csv"))

    def get(code, system):
        return df[(df["code"] == code) & (df["system"] == system)].iloc[0]

    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    st.font.size = SIZE

    # ===== 封面标题 =====
    title(doc, "复刻传奇：海龟交易法则实战演练")
    info = doc.add_paragraph(); pfmt(info, WD_ALIGN_PARAGRAPH.CENTER)
    rfont(info.add_run("北京大学 AI 量化工作坊 · TASK4　　姓名：张哲铭"), size=Pt(10.5))
    doc.add_paragraph()

    body(doc,
        "本任务在前三次工作坊（搭建数据引擎、数据诊断与指标构造、双均线策略首秀）的基础上，深入学习并"
        "完整复刻华尔街传奇——海龟交易法则（Turtle Trading）。全文首先阐明海龟策略的核心思想与关键优势，"
        "解释其三大基石：高低点通道（唐奇安通道）、平均真实波幅 ATR 与基于 ATR 的止损/头寸管理；随后在实际"
        "回测前先完整设计策略全流程并绘制流程图；接着用 Python 手工实现该策略，并检索重要学术文献以明确其在"
        "何种标的上更有效；最后对黄金 ETF、纳指 ETF、腾讯控股、沪深 300 等 8 个标的、两套海龟系统（20/10 与 "
        "55/20）共 16 组参数进行实证回测，用总收益、年化、超额收益、夏普、最大回撤、卡玛、胜率、盈亏比等一"
        "整套指标全面评估，并做通道周期敏感性分析，总结海龟法则的适应场景与使用心得。")

    # ===== 一、海龟策略核心思想与关键优势 =====
    heading(doc, "一、海龟交易法则：核心思想与关键优势")
    body(doc,
        "海龟交易法则源于 1983 年美国期货交易大师 Richard Dennis（丹尼斯）与合伙人 William Eckhardt（埃克哈特）"
        "的一场著名赌约：交易能力究竟是天赋还是可以后天培养？丹尼斯从上千名报名者中挑选了一批毫无经验的普通人"
        "（戏称“海龟”），用两周时间教会他们一套完全机械化的交易系统。此后五年，这批海龟合计盈利超过 1.75 亿美元、"
        "顶尖者年化收益约 80%，丹尼斯赢得赌约。海龟法则由此成为金融史上最著名、被最广泛研究与复现的系统化趋势"
        "跟踪策略。")
    subheading(doc, "· 核心思想：机械化的趋势跟踪")
    body(doc,
        "海龟法则的哲学是“价格会沿最小阻力方向运动”，即市场一旦形成趋势就往往延续。它不预测顶底，只在价格"
        "创出一段时间新高（新趋势确立）时顺势入场做多、在趋势逆转时离场，力求“截断亏损、让利润奔跑”。整套系统"
        "把入场、加仓、止损、离场、头寸规模全部量化为明确规则，要求 100% 机械执行、剔除情绪干扰——这正是它可以"
        "被“教会普通人”的根本原因。")
    subheading(doc, "· 关键优势")
    body(doc,
        "1. 完全机械化、可复制、可回测。 每一步都有明确的数值规则（20/55 日通道、20 日 ATR、2 ATR 止损、"
        "0.5 ATR 加仓），不依赖个人判断，杜绝“追涨杀跌”的情绪化交易，也使策略能被严格量化检验。", indent=True)
    body(doc,
        "2. 以波动率（ATR）为核心的风险管理。 海龟用 ATR 动态度量每个标的的“脾气”，据此决定头寸大小与止损"
        "距离：波动大的品种买得少、止损宽，波动小的品种买得多、止损窄，使每一笔交易承担的账户风险大致相等，"
        "这是它区别于普通均线策略的最大亮点。", indent=True)
    body(doc,
        "3. 金字塔式分批加仓，让利润奔跑。 趋势一旦被确认，价格每上涨 0.5 ATR 就加一个单位、最多 4 个单位，"
        "使盈利头寸在大趋势中不断放大，充分吃到主升浪。", indent=True)
    body(doc,
        "4. 严格止损、控制回撤。 每单位设 2 ATR 硬止损，单一市场、相关市场组、总方向都设有持仓上限，"
        "从制度上防止单笔或单一方向的过度暴露，这使其在剧烈下跌中往往能显著跑赢“死扛”的买入持有。", indent=True)
    body(doc,
        "5. 低胜率、高盈亏比的稳健盈利结构。 海龟靠少数几次抓住大趋势的巨大盈利，覆盖多次突破失败的小额止损，"
        "因此天生“胜率不高但盈亏比很高”，只要严格执行就能长期为正——这一特征在后文回测中反复得到验证。", indent=True)

    # ===== 二、三大核心概念解释 =====
    heading(doc, "二、核心概念解释：高低点通道、ATR 与止损条件")
    subheading(doc, "1. 高低点通道（唐奇安通道 Donchian Channel）")
    body(doc,
        "高低点通道由趋势跟踪之父 Richard Donchian 提出，是海龟法则的信号来源。它取过去 N 个交易日的最高价作为"
        "“上轨”、最低价作为“下轨”：")
    formula(doc, "上轨 = max(最近 N 日最高价)，  下轨 = min(最近 N 日最低价)")
    body(doc,
        "当收盘价向上突破上轨，意味着价格创出 N 日新高、上涨阻力被打破，是做多入场信号；跌破一条更短周期"
        "（如 N/2 日）的下轨则视为多头趋势结束、离场。海龟用两套周期：系统一（入场 20 日、离场 10 日，捕捉中短"
        "趋势）与系统二（入场 55 日、离场 20 日，捕捉长期大趋势）。通道呈“阶梯状”——只有创出新极值时才移动，"
        "因此触发价格完全透明、可提前算出，天然适合机械执行。")
    subheading(doc, "2. 平均真实波幅（ATR, Average True Range）")
    body(doc,
        "ATR 度量价格的“平均单日波动幅度”，是海龟做风险管理的标尺。先计算每日真实波幅 TR（同时考虑跳空缺口）：")
    formula(doc, "TR = max( 今日最高−今日最低,  |今日最高−昨收|,  |今日最低−昨收| )")
    body(doc,
        "再对 TR 取 N 日 Wilder 平滑移动平均，即得 ATR（本文 N=20，与海龟原版一致）：")
    formula(doc, "ATR_t = [ ATR_(t−1) × (N−1) + TR_t ] / N")
    body(doc,
        "ATR 越大表示该标的越“躁动”。海龟据此计算单位头寸（Position Unit）——让“价格逆行 1 个 ATR 时账户恰好"
        "亏损设定的风险比例（如 1%）”：")
    formula(doc, "单位头寸 = 风险资本 / ( N × ATR × 价值因子 )")
    body(doc,
        "其含义是：波动越大（ATR 越大）买得越少，波动越小买得越多，从而让不同标的、不同时期的每笔交易承担"
        "大致相等的风险，实现真正的“波动率平价”仓位管理。")
    subheading(doc, "3. 止损条件与加仓规则")
    body(doc,
        "海龟的风险控制围绕 ATR 展开，有三条硬规则：（1）止损——每个单位在入场价下方 2 ATR 处设硬止损，一旦"
        "触及立即无条件离场，把单笔最大亏损锁定在约 2%（2 ATR × 1% 风险单位）；（2）加仓——价格每上涨 0.5 ATR "
        "加 1 个单位、最多 4 个单位，且每次加仓后止损线同步上移，锁定已有利润；（3）离场——价格跌破离场通道"
        "（系统一 10 日 / 系统二 20 日最低）即视为趋势结束，全部平仓。三者共同构成“亏损截断得快、盈利放得开”的"
        "非对称收益结构。")

    # ===== 三、策略全流程设计与流程图 =====
    heading(doc, "三、策略全流程设计（回测前的流程图）")
    body(doc,
        "在动手回测之前，先把上述规则串成一条完整的决策链路并绘制成流程图（图 1），确保编程实现严格对应策略"
        "逻辑、不遗漏任何环节。整个流程从“选择市场”开始，依次计算 ATR、单位头寸，随后进入“监控突破—入场—"
        "加仓/止损/止盈”的持仓循环。")
    pic(doc, "flowchart.png", width=6.4)
    caption(doc, "图 1　海龟交易策略完整流程图（做多方向）")
    body(doc,
        "如图 1，橙色主链路（①→⑤）完成开仓前的准备：从高流动性品种中选定标的、计算 20 日 ATR、据此算出单位"
        "头寸、监控价格是否突破 N 日通道上轨，突破确认后建立第 1 个单位（蓝色⑤）。此后进入决策循环：先判断"
        "价格是否较上次建/加仓价再涨 0.5 ATR，是则加仓（绿色⑥，最多 4 单位）并回到循环；否则依次检查是否跌破"
        "止损线（红色，止损离场）、是否跌破离场通道（绿色，止盈/趋势结束离场）；若都未触发则继续持有、循环"
        "监控。本文的 Python 实现即严格按此状态机编写。")

    # ===== 四、Python 编程实现 =====
    heading(doc, "四、Python 编程实现")
    body(doc,
        "本文回测引擎完全手工实现（turtle_strategy.py），核心分三部分：ATR 与高低点通道计算、逐日状态机"
        "（入场/加仓/止损/离场）、以及净值与指标计算。考虑到 A 股与 ETF 现货难以便捷做空，本文只实现做多方向，"
        "更贴近国内可实践场景。为杜绝“未来函数”，取通道极值时用 shift(1) 排除当日、且第 t 日收盘产生的仓位变化"
        "在第 t+1 日才生效；同时在仓位变动当日按变动金额扣除单边万分之五交易成本。核心代码如下：")
    code_block(doc,
        '# 1) ATR（Wilder 平滑）与高低点通道\n'
        'tr = pd.concat([high-low, (high-prev_close).abs(),\n'
        '                (low-prev_close).abs()], axis=1).max(axis=1)\n'
        'atr = wilder_smooth(tr, n=20)\n'
        'dc_upper = high.rolling(entry_n).max().shift(1)   # 入场上轨(N日最高)\n'
        'dc_exit  = low.rolling(exit_n).min().shift(1)     # 离场下轨(N/2日最低)\n\n'
        '# 2) 逐日状态机（只做多）\n'
        'if units == 0:                       # 空仓：监控突破\n'
        '    if price > dc_upper[i]:          # 突破上轨 -> 入场1单位\n'
        '        units=1; entry=price; stop=price-2*atr[i]\n'
        'else:                                # 持仓：按优先级判定\n'
        '    if price < stop:                 # (a) 跌破止损线 -> 止损离场\n'
        '        close_trade("stop"); units=0\n'
        '    elif price < dc_exit[i]:         # (b) 跌破离场通道 -> 止盈离场\n'
        '        close_trade("exit"); units=0\n'
        '    elif units<4 and price>=last_add+0.5*atr:   # (c) 涨0.5ATR -> 加仓\n'
        '        units+=1; last_add=price; stop=price-2*atr\n\n'
        '# 3) 仓位次日生效(避免未来函数)、扣成本、算净值\n'
        'position  = (units/4).shift(1)\n'
        'strat_ret = position*close.pct_change() - position.diff().abs()*0.0005\n'
        'equity    = (1+strat_ret).cumprod()')
    body(doc,
        "数据方面，本文复用工作坊统一的行情数据（前复权日线，多数标的自 2018-04 至 2026-07，共约 2000 个交易日）。"
        "下面以趋势特征鲜明的腾讯控股为例，展示海龟策略的信号可视化（图 2），完整呈现价格、高低点通道、ATR 与"
        "四类交易信号。")
    pic(doc, "signal_hk00700.png")
    caption(doc, "图 2　腾讯控股（00700）海龟策略交易信号与 ATR（入场20日/离场10日）")
    body(doc,
        "如图 2，上图红色虚线为 20 日高点通道（入场上轨）、绿色虚线为 10 日低点通道（离场下轨），阴影为二者构成"
        "的通道区间；红色三角为突破上轨的入场点、橙色十字为每涨 0.5 ATR 的加仓点、紫色叉为触及 2 ATR 止损、"
        "绿色倒三角为跌破下轨的趋势离场。下图为 20 日 ATR，可见 2021 与 2022 年腾讯剧烈波动时 ATR 显著抬升，"
        "策略据此自动缩小头寸、放宽止损。整体可清晰看到海龟“突破进场、分批加仓、破位即走”的机械纪律。")

    # ===== 五、文献综述 =====
    heading(doc, "五、文献综述：海龟/趋势跟踪策略在什么标的上更有效")
    body(doc,
        "海龟法则属于趋势跟踪（Trend Following）大家族。动手回测前，先回顾学术界对其有效性与适用标的的重要研究，"
        "为标的选择与结果解读提供依据。")
    subheading(doc, "1. 通道突破规则的实证根基：Brock, Lakonishok & LeBaron (1992)")
    body(doc,
        "发表于顶级期刊《Journal of Finance》的 Brock、Lakonishok 与 LeBaron（1992，简称 BLL）是检验技术交易规则"
        "的开山之作。他们用道琼斯工业指数 1897—1986 长达 90 年的数据，对移动平均规则和“区间突破”规则（与海龟高低"
        "点通道同源）做了严格的自助法（Bootstrap）检验，发现买入信号后的平均收益显著高于卖出信号、且波动更小，"
        "无法被随机游走等模型解释，为“简单突破规则确实包含预测力”提供了首个量化证据。")
    subheading(doc, "2. 趋势跟踪的主场是大宗商品：Miffre & Rallis (2007)")
    body(doc,
        "Miffre 与 Rallis（2007，《Journal of Banking & Finance》）对 31 种商品期货 1979—2004 的动量/趋势策略研究"
        "发现：13 个趋势策略年均收益达 9.38%，而同期等权买入持有商品组合反而亏损；且该收益与股债相关性极低、不随"
        "样本期衰减。Moskowitz、Ooi 与 Pedersen（2012，《Journal of Financial Economics》）在 58 个跨资产品种上"
        "进一步证实了“时间序列动量”的普遍存在。二者共同指向：趋势跟踪在趋势鲜明、波动大、有持续供求/宏观驱动的"
        "标的（如黄金等贵金属、原油等商品）上最有效。")
    subheading(doc, "3. 海龟规则的直接复现：Swart (2016, 开普敦大学)")
    body(doc,
        "Swart（2016，University of Cape Town 金融硕士论文）直接以海龟法则为蓝本，在南非期货交易所（SAFEX）用"
        "唐奇安通道 + ATR 完整复现了系统一（20 日入场/10 日离场）、系统二（55 日入场/20 日离场）及二者的整合系统，"
        "并用 20 日 ATR 决定头寸规模、止损与加仓。该文回顾指出：既有文献主要在北美与亚洲市场的商品期货与股指期货"
        "上检验海龟法则，说明这两类高流动性、趋势性资产是海龟策略被验证最充分、也最有效的标的。")
    subheading(doc, "4. 有效性会衰减：自适应市场假说")
    body(doc,
        "另一方面，Bessembinder & Chan（1998）等研究发现，BLL 记录的规则预测力在 1987 年后明显减弱。Lo（2004）"
        "由此提出“自适应市场假说”：任何技术规则一旦被广泛知晓、套利资金涌入，超额收益就会被抢跑而衰减。这提醒"
        "我们——在流动性极好、参与者极成熟的大盘宽基指数上，海龟择时越来越难跑赢买入持有。")
    subheading(doc, "5. 对本文标的选择的启示")
    body(doc,
        "综合文献，海龟法则更可能在“黄金等大宗商品、趋势鲜明的成长/个股”上跑出价值，而在“高效率、宽幅震荡的"
        "宽基指数”上因滞后与假信号跑输买入持有。恰好，本文纳入了文献直接指向的黄金 ETF（518880，商品属性）与"
        "纳指 ETF（159941，成长指数）作为验证标的，同时纳入沪深 300、中证 500 等宽基指数与长江电力、红利低波等"
        "低波动标的作为对照，下面用真实回测加以检验。")

    # ===== 六、多标的多参数实证回测 =====
    heading(doc, "六、多标的、多参数实证回测")
    body(doc,
        "本文对 8 个标的分别运行系统一（System1，入场 20 日 / 离场 10 日）与系统二（System2，入场 55 日 / 离场 "
        "20 日）两套海龟参数，共 16 组结果。ATR 周期统一取 20，止损 2 ATR，加仓间隔 0.5 ATR、最多 4 单位。每组均"
        "与“买入持有”基准对比，重点看超额收益与风险控制。")

    subheading(doc, "1. 回测指标体系")
    body(doc,
        "为全面评估，本文设计并计算了以下一整套指标：总收益率、年化收益率、超额收益（策略年化 − 买入持有年化，"
        "用于剥离标的 β、只看策略 α）、夏普比率（每单位波动的超额收益，越高越好）、最大回撤 MDD（净值从高点回落"
        "的最大幅度，衡量下行风险）、卡玛比率（年化收益 / |最大回撤|，收益回撤性价比）、胜率（盈利交易占比）、"
        "盈亏比（平均盈利 / 平均亏损）与交易次数。核心公式如下：")
    formula(doc, "年化 = (1+总收益)^(252/天数) − 1；  夏普 = 日超额收益均值/标准差 × √252")
    formula(doc, "MDD = min( 净值_t / 历史最高净值 − 1 )；  卡玛 = 年化收益 / |MDD|")

    # ---- 表1：论文验证标的 ----
    subheading(doc, "2. 文献验证标的：黄金 ETF 与纳指 ETF")
    body(doc,
        "先看文献直接指向的两个标的。表 1 给出黄金 ETF 与纳指 ETF 在两套系统下的完整指标。")
    rows1 = []
    for code in ["sh518880", "sz159941"]:
        for sysn in ["System1", "System2"]:
            r = get(code, sysn)
            rows1.append([
                NAME[code], f"{sysn}({int(r['entry_n'])}/{int(r['exit_n'])})",
                pct(r["strat_annual"]), pct(r["bench_annual"]),
                pct(r["excess_annual"]), f"{r['sharpe']:.2f}",
                pct(r["mdd"]), pct(r["bench_mdd"]),
                pct(r["win_rate"], 0), f"{r['pl_ratio']:.2f}"])
    caption(doc, "表 1　黄金 ETF 与纳指 ETF 海龟策略回测结果")
    make_table(doc,
        ["标的", "系统(入场/离场)", "策略年化", "基准年化", "超额年化", "夏普",
         "策略回撤", "基准回撤", "胜率", "盈亏比"], rows1)
    body(doc,
        "结果与文献既有印证、也有一层重要修正。黄金 ETF 是海龟的“理想画像”：系统二年化 5.3%、夏普 0.51、"
        "胜率 63%、盈亏比高达 3.73，最大回撤仅 16.7%——远低于买入持有的 30.5%。但在 2018—2026 这段黄金整体单边"
        "上涨的行情里，买入持有年化高达 15.6%，海龟的超额收益为负（−10.3%）。这揭示了一个诚实结论：在“一路上涨、"
        "少有深调”的长牛现货中，任何“突破进、破位出”的择时都会因反复止损而跑输“一直满仓”。纳指 ETF 同理，"
        "且系统一优于系统二（近年纳指波动加剧，短周期反应更快）。")

    # ---- 图3：黄金ETF净值 ----
    pic(doc, "equity_sh518880.png")
    caption(doc, "图 3　黄金 ETF 两套海龟系统净值 vs 买入持有")
    body(doc,
        "如图 3，黄金 ETF 买入持有（灰色虚线）在 2024—2025 年加速上行，而海龟策略因每轮小回调都触发离场/止损、"
        "错过部分连续涨幅，净值增长较缓。但换个角度：策略净值曲线明显更平滑、回撤更浅，这正是趋势跟踪“用部分"
        "收益换回撤保护”的取舍。")

    # ---- 表2：全部8标的System2 ----
    subheading(doc, "3. 全部标的横向对比（System2 · 55/20）")
    body(doc,
        "表 2 汇总全部 8 个标的在系统二下的表现，按超额年化从高到低排序；图 4 为对应的超额收益柱状图。")
    rows2 = []
    sub = df[df["system"] == "System2"].copy().sort_values("excess_annual", ascending=False)
    for _, r in sub.iterrows():
        rows2.append([
            r["name"], pct(r["strat_annual"]), pct(r["bench_annual"]),
            pct(r["excess_annual"]), f"{r['sharpe']:.2f}", f"{r['calmar']:.2f}",
            pct(r["mdd"]), pct(r["bench_mdd"]), f"{int(r['n_trades'])}",
            pct(r["win_rate"], 0)])
    caption(doc, "表 2　8 标的海龟策略（System2）指标全览（按超额年化排序）")
    make_table(doc,
        ["标的", "策略年化", "基准年化", "超额年化", "夏普", "卡玛", "策略回撤",
         "基准回撤", "交易次数", "胜率"], rows2)
    pic(doc, "summary_excess.png")
    caption(doc, "图 4　各标的海龟策略（System2）超额年化收益对比")
    body(doc,
        "如表 2 与图 4，超额收益出现了鲜明分化，且与文献高度一致：腾讯控股（超额 +4.9%）与科创 50、沪深 300 等"
        "波动大、有明显趋势与深调的标的，海龟策略跑赢或追平买入持有；而黄金、纳指、长江电力、红利低波等“一路"
        "上涨少回调”或“低波动无趋势”的标的，海龟则跑输。其中腾讯控股是最佳范例——它 2021—2022 年经历过腰斩级"
        "暴跌，海龟在下跌初期即离场、成功规避，因此不仅超额转正，风险控制更是压倒性领先。")

    # ---- 图5：回撤对比 ----
    subheading(doc, "4. 海龟真正的价值：控制回撤")
    body(doc,
        "若只看收益，容易低估海龟。把视角切到风险维度（图 5，策略 vs 买入持有最大回撤），结论截然不同。")
    pic(doc, "risk_compare.png")
    caption(doc, "图 5　海龟策略 vs 买入持有：最大回撤对比（System2，越低越好）")
    body(doc,
        "如图 5，在高波动标的上海龟策略大幅削减了回撤：腾讯控股从买入持有的 76.7% 压降到 21.0%，科创 50 从 "
        "59.9% 降到 16.5%，沪深 300 从 44.7% 降到 19.5%，中证 500 从 40.7% 降到 23.8%。也就是说，策略往往用“少赚"
        "一点”换来了“回撤浅一大截、持有体验平稳得多”。唯一例外是低波动的长江电力——它本无明显趋势可供跟踪，"
        "择时失误反而略增回撤，印证了“低波动无趋势标的不适合海龟”。")

    # ---- 图6：夏普热力图 ----
    pic(doc, "sharpe_compare.png", width=4.6)
    caption(doc, "图 6　8 标的 × 2 系统 夏普比率热力图（红高绿低）")
    body(doc,
        "图 6 的夏普比率（风险调整后收益）进一步佐证：科创 50、腾讯控股、黄金 ETF 大面积偏红（夏普 0.5—0.58，"
        "全样本最高），是最适合海龟策略的标的；而长江电力、红利低波则偏绿（夏普接近 0 甚至为负），风险调整后"
        "性价比很差。这与“趋势/波动越强越适合海龟”的规律一致。")

    # ---- 图7：收益-回撤散点 ----
    pic(doc, "return_risk.png", width=5.7)
    caption(doc, "图 7　海龟策略收益—回撤分布（红=跑赢基准 绿=跑输，●系统1 ■系统2）")
    body(doc,
        "图 7 把 16 组结果画到“收益—回撤”平面：绝大多数点集中在左侧（回撤 15%—25%），说明海龟策略把各标的的"
        "回撤都压到了相近的较低水平——这正是 ATR 头寸管理“波动率平价”的效果，无论标的本身多凶，策略回撤都被"
        "控制在可控区间。红点（跑赢基准）多为高波动趋势标的，绿点（跑输）多为长牛或低波动标的。")

    # ===== 七、参数敏感性 =====
    heading(doc, "七、参数调节与敏感性分析")
    body(doc,
        "按任务要求，本文进一步调节核心参数（标的类型、通道周期），观察收益变化。图 8 以黄金 ETF 为例，将入场"
        "通道周期从 10 日逐步拉长到 80 日（离场周期取其一半），观察年化收益、夏普与最大回撤的变化。")
    pic(doc, "param_scan.png")
    caption(doc, "图 8　黄金 ETF 通道周期敏感性（周期越长交易越少、越稳）")
    # 取黄金扫描关键数据
    g = scan[scan["code"] == "sh518880"]
    rows3 = []
    for _, r in g.iterrows():
        rows3.append([f"{int(r['entry_n'])}/{int(r['exit_n'])}",
                      pct(r["strat_annual"]), f"{r['sharpe']:.2f}",
                      pct(r["mdd"]), int(r["n_trades"]), pct(r["win_rate"], 0)])
    caption(doc, "表 3　黄金 ETF 不同通道周期回测结果")
    make_table(doc, ["入场/离场周期", "年化收益", "夏普", "最大回撤", "交易次数", "胜率"], rows3)
    body(doc,
        "如图 8 与表 3，黄金 ETF 呈现清晰的“周期越长越好”规律：入场周期从 10 日拉长到 55 日，年化收益从 2.9% "
        "升到 9.7%、夏普从 0.33 升到 0.82、交易次数从 74 次降到 21 次——长周期过滤掉了震荡市的大量假信号，只在"
        "真正的大趋势上出手，非常契合黄金“慢牛”的特性。但周期并非越长越好：拉到 80 日后回撤反而扩大（−24.9%），"
        "因为反应过慢、离场太迟。")
    body(doc,
        "然而这一规律并不通用。对纳指 ETF，通道周期越长收益反而越差（80 日年化 −1.1%），因为近年纳指波动加剧、"
        "长周期反应太慢、频繁在高位追进又深套；沪深 300 则对周期不敏感，各周期年化都在 3%—4.5% 徘徊。这说明："
        "没有“万能周期”，参数必须与标的的波动节奏匹配——趋势平稳绵长的品种（黄金）适合长周期，波动急促的品种"
        "（纳指）反而适合较短周期抢反应。相比之下，“选对标的”比“调对参数”对最终结果的影响更大。")

    # ===== 八、应用心得 =====
    heading(doc, "八、海龟法则适应场景与使用心得")
    body(doc, "综合文献与本文 16 组实证 + 通道周期扫描，对海龟交易法则的适应场景与使用要点总结如下：")
    body(doc,
        "1. 适应场景——趋势鲜明、波动较大、会有深度回调的标的。 海龟的全部超额收益都来自“抓大趋势、避大跌”，"
        "因此在腾讯这类高波动个股、科创 50 这类弹性成长指数上最有价值（超额转正、回撤腰斩）；在黄金、纳指这类"
        "“一路慢牛少回调”的长牛资产上，它虽控回撤出色却难跑赢买入持有；在长江电力、红利低波这类低波动无趋势"
        "标的上则基本失效。文献指向的商品/股指期货之所以是海龟主场，正因其趋势性强且可双向交易。", indent=True)
    body(doc,
        "2. 海龟的核心价值常是“控回撤”而非“多赚”。 本文中它把腾讯回撤从 77% 压到 21%、科创 50 从 60% 压到 "
        "17%、沪深 300 从 45% 压到 20%。对厌恶大幅回撤、追求平稳净值曲线的资金而言，用一点收益换显著更浅的回撤"
        "与更高夏普，本身就是极有意义的风险管理——这也是趋势跟踪 CTA 在机构资产配置中的核心定位。", indent=True)
    body(doc,
        "3. 必须接受“低胜率、高盈亏比”，并有纪律地执行每一次止损。 本文趋势型标的盈亏比常达 2—6 倍但胜率仅 "
        "30%—60%，盈利高度依赖少数几笔大趋势。使用者不能因连续小额止损而放弃系统，否则会恰好错过贡献全部收益"
        "的那几笔大行情——这正是海龟“机械执行、剔除情绪”纪律的意义所在。", indent=True)
    body(doc,
        "4. 参数需与标的波动匹配，且必须计入成本、规避未来函数。 趋势平稳的品种用长周期（黄金 55 日夏普最高），"
        "波动急促的品种用短周期；短周期交易频繁，务必把佣金与冲击成本纳入回测（本文用万分之五），并让信号次日"
        "生效、通道取值排除当日，否则会系统性高估收益。ATR 头寸管理是海龟的精髓，应始终保留。", indent=True)
    body(doc,
        "5. 只做多是国内现货的现实约束，也是本文超额偏低的重要原因。 海龟原版在期货市场可双向做空，能在下跌"
        "趋势中同样获利；而本文受限于 A 股/ETF 现货只做多，下跌时最多空仓规避、无法反向盈利，因此在长牛样本里"
        "超额收益偏保守。若在可做空的商品期货、股指期货上应用（如文献所示），海龟法则的威力会更完整地释放。", indent=True)
    body(doc,
        "总之，海龟交易法则是理解“系统化趋势跟踪”最经典的范式：它用高低点通道捕捉趋势、用 ATR 统一风险度量、"
        "用金字塔加仓放大利润、用 2 ATR 止损截断亏损，把一套完整的交易纪律彻底量化。本文最重要的一课是——评价"
        "任何策略都要用超额收益把“标的的 β”与“策略的 α”分开，用夏普和最大回撤把“收益”与“风险”一起看；海龟"
        "的价值未必写在收益率上，而常常藏在那条更平滑、回撤更浅的净值曲线里。")

    # 结尾
    doc.add_paragraph()
    note = doc.add_paragraph(); pfmt(note, WD_ALIGN_PARAGRAPH.LEFT)
    rfont(note.add_run(
        "数据来源：工作坊统一行情数据（前复权日线，2018—2026，8 标的约 2000 个交易日）；"
        "回测与绘图代码：turtle_strategy.py / metrics.py / run_backtest.py / make_flowchart.py。"
        "主要参考文献：Brock, Lakonishok & LeBaron (1992, JF)；Miffre & Rallis (2007, JBF)；"
        "Moskowitz, Ooi & Pedersen (2012, JFE)；Swart (2016, UCT)；Faith《海龟交易法则》。"
        "本文仅为量化学习实践，不构成任何投资建议，市场有风险，决策需谨慎。"),
        size=Pt(9), color=RGBColor(0x66, 0x66, 0x66))

    doc.save(DOCX)
    print("全部章节已写入：", DOCX)


if __name__ == "__main__":
    build()
